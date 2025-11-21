# formatter-service/formatting/abbreviations.py

import re
from typing import List, Tuple
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK

# Pattern for lines like:
#   GDP – Gross Domestic Product
#   IMF - International Monetary Fund
#   UB: University of Buea
ABBR_LINE_REGEX = re.compile(
    r"^\s*([A-Z]{2,10})\s*[-–—:]\s+(.+)$"
)


def collect_abbreviations(doc: Document) -> List[Tuple[str, str]]:
    """
    Scan the entire document for lines that look like:
      'GDP – Gross Domestic Product'
      'IMF - International Monetary Fund'
    and return a list of (abbr, full) pairs (deduplicated, sorted).
    """
    found = {}
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue

        m = ABBR_LINE_REGEX.match(text)
        if not m:
            continue

        abbr = m.group(1).strip()
        full = m.group(2).strip()

        if len(abbr) < 2 or len(full) < 3:
            continue

        if abbr not in found:
            found[abbr] = full

    return sorted(found.items(), key=lambda x: x[0])


def _has_abbreviations_heading(doc: Document) -> bool:
    """Avoid inserting ABBREVIATIONS twice."""
    for p in doc.paragraphs:
        if (p.text or "").strip().upper() == "ABBREVIATIONS":
            return True
    return False


def _find_anchor_for_abbreviations(doc: Document):
    """
    Find where ABBREVIATIONS should come after.
    Priority:
      1) LIST OF FIGURES
      2) LIST OF TABLES
      3) TABLE OF CONTENTS
    """
    lof = None
    lot = None
    toc = None

    for p in doc.paragraphs:
        text = (p.text or "").strip().upper()
        if text == "LIST OF FIGURES":
            lof = p
        elif text == "LIST OF TABLES":
            lot = p
        elif text == "TABLE OF CONTENTS":
            toc = p

    return lof or lot or toc


def insert_abbreviations_section(doc: Document, abbrevs: List[Tuple[str, str]]) -> None:
    """
    Insert an 'ABBREVIATIONS' heading + 2-column table on its own page,
    AFTER (in order of preference): LIST OF FIGURES, LIST OF TABLES, or TOC.
    """
    if not abbrevs:
        return

    if _has_abbreviations_heading(doc):
        return

    anchor_p = _find_anchor_for_abbreviations(doc)

    # If absolutely no anchor, we put it near the top
    if anchor_p is None:
        if not doc.paragraphs:
            anchor_el = None
        else:
            anchor_el = doc.paragraphs[0]._p
    else:
        anchor_el = anchor_p._p

    # Create ABBREVIATIONS heading and table at the end
    abbr_title_p = doc.add_paragraph("ABBREVIATIONS")
    try:
        abbr_title_p.style = doc.styles["Heading 1"]
    except Exception:
        pass

    pf = abbr_title_p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before = Pt(24)
    pf.space_after = Pt(12)

    for run in abbr_title_p.runs:
        font = run.font
        font.bold = True
        font.size = Pt(14)

    # Create table (initially appended at end)
    table = doc.add_table(rows=len(abbrevs) + 1, cols=2)

    try:
        table.style = "Table Grid"
    except Exception:
        pass

    # Header row
    hdr = table.rows[0].cells
    hdr[0].text = "Abbreviation"
    hdr[1].text = "Meaning"
    for cell in hdr:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(11)

    # Data rows
    for i, (abbr, full) in enumerate(abbrevs, start=1):
        row = table.rows[i].cells
        row[0].text = abbr
        row[1].text = full
        for cell in row:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(11)

    # XML elements
    abbr_el = abbr_title_p._p
    tbl_el = table._tbl

    # Page break before ABBREVIATIONS page
    pb_para = doc.add_paragraph()
    pb_run = pb_para.add_run()
    pb_run.add_break(WD_BREAK.PAGE)
    pb_el = pb_para._p

    if anchor_el is None:
        # No anchor: drop ABBREVIATIONS near top
        if doc.paragraphs:
            first_el = doc.paragraphs[0]._p
            first_el.addprevious(pb_el)
            pb_el.addnext(abbr_el)
            abbr_el.addnext(tbl_el)
        else:
            # Edge case: empty doc, leave as is
            return
    else:
        # Normal case: insert after anchor (LOT or LOF or TOC)
        anchor_el.addnext(pb_el)
        pb_el.addnext(abbr_el)
        abbr_el.addnext(tbl_el)
