# formatter-service/formatting/abbreviations_builder.py

"""
LIST OF ABBREVIATIONS generation.

Scans body text for acronyms (e.g. "GDP", "HIV/AIDS", "ICT") and, where the
document defines them in the usual academic way — "Information and
Communication Technology (ICT)" — captures the expansion too. The collected
entries are rendered as a prelim page with one tab-aligned line per entry.
"""

from __future__ import annotations

import re
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.shared import Cm
from docx.text.paragraph import Paragraph

# Words that look like acronyms but are just shouting or structure
_STOPWORDS = {
    "CHAPTER", "TABLE", "FIGURE", "LIST", "OF", "THE", "AND", "FOR",
    "ABSTRACT", "REFERENCES", "BIBLIOGRAPHY", "APPENDIX", "APPENDICES",
    "CONTENTS", "ABBREVIATIONS", "ACKNOWLEDGEMENT", "ACKNOWLEDGEMENTS",
    "DEDICATION", "INTRODUCTION", "CONCLUSION", "SUMMARY", "PAGE",
    "II", "III", "IV", "VI", "VII", "VIII", "IX", "XI", "XII",
}

_ACRONYM_RE = re.compile(r"\b[A-Z][A-Z0-9]{1,9}(?:/[A-Z0-9]{2,9})?\b")

# "Some Expanded Phrase (ABBR)" — words immediately before a parenthesised acronym
_DEFINITION_RE = re.compile(
    r"((?:[A-Za-z][\w'’-]*\s+){1,9}[A-Za-z][\w'’-]*)\s*\(\s*([A-Z][A-Z0-9]{1,9})s?\s*\)"
)


def extract_abbreviations(
    blocks: List[Dict[str, Any]],
    metadata: Dict[str, Any],
) -> List[Tuple[str, str]]:
    """
    Return a sorted list of (abbreviation, expansion) tuples found in the
    document's paragraph blocks. Expansion is "" when we can't infer one.
    """
    block_meta: Dict[str, Any] = metadata.get("blocks", {}) or {}

    found: Dict[str, str] = {}

    for block in blocks:
        if block.get("type") != "paragraph":
            continue

        text = (block.get("text") or "").strip()
        if not text:
            continue

        # Skip headings/captions — acronyms there are usually titles
        role = (block_meta.get(block.get("id"), {}) or {}).get("role", "")
        if role in (
            "chapter_heading", "section_heading", "subsection_heading",
            "table_of_contents", "list_of_tables", "list_of_figures",
            "title_page",
        ):
            continue

        # 1) Definitions: "Expanded Phrase (ABBR)"
        for match in _DEFINITION_RE.finditer(text):
            phrase, abbr = match.group(1).strip(), match.group(2)
            if _is_stopword(abbr):
                continue
            expansion = _trim_phrase_to_acronym(phrase, abbr)
            if expansion:
                found[abbr] = expansion
            elif abbr not in found:
                found[abbr] = ""

        # 2) Bare acronyms
        for match in _ACRONYM_RE.finditer(text):
            abbr = match.group(0)
            base = abbr.split("/")[0]
            if _is_stopword(base) or len(base) < 2:
                continue
            if abbr not in found:
                found[abbr] = ""

    return sorted(found.items(), key=lambda kv: kv[0])


_CONNECTIVES = {"of", "and", "the", "for", "in", "on", "to", "a", "an"}


def _is_stopword(token: str) -> bool:
    """Stopword check that also catches simple plurals (TABLE/TABLES)."""
    if token in _STOPWORDS:
        return True
    return token.endswith("S") and token[:-1] in _STOPWORDS


def _trim_phrase_to_acronym(phrase: str, abbr: str) -> str:
    """
    Given the words captured before "(ABBR)", return the shortest trailing
    sub-phrase whose word initials spell the acronym (connectives ignored),
    e.g. "and the Gross Domestic Product" + GDP -> "Gross Domestic Product".
    Returns "" when no sub-phrase matches.
    """
    words = [w for w in re.split(r"[\s-]+", phrase) if w]
    target = abbr.upper()

    for start in range(len(words) - 1, -1, -1):
        candidate = words[start:]
        if candidate[0].lower() in _CONNECTIVES:
            continue
        significant = [w for w in candidate if w.lower() not in _CONNECTIVES]
        initials = "".join(w[0].upper() for w in significant)
        if initials == target or "".join(w[0].upper() for w in candidate) == target:
            return " ".join(candidate)

    return ""


def insert_list_of_abbreviations(
    doc: Document,
    abbreviations: List[Tuple[str, str]],
    anchor_para: Optional[Paragraph] = None,
    title_style_name: str = "Heading 1",
) -> Optional[Paragraph]:
    """
    Insert a LIST OF ABBREVIATIONS page BEFORE the anchor paragraph,
    preceded by a page break. Each entry is a tab-aligned line:

        ABBR<tab>Expansion

    Returns the inserted title paragraph (the new anchor), or None if
    there was nothing to insert.
    """
    if not abbreviations:
        return None

    # --- 1. PAGE BREAK ---
    is_first_paragraph_in_doc = (
        anchor_para is not None
        and doc.paragraphs
        and anchor_para._p is doc.paragraphs[0]._p
    )

    break_para = None
    if anchor_para is not None and not is_first_paragraph_in_doc:
        break_para = anchor_para.insert_paragraph_before()
        break_para.add_run().add_break(WD_BREAK.PAGE)

    # --- 2. ENTRY LINES (inserted before anchor, so build top-down) ---
    def _new_para() -> Paragraph:
        if anchor_para is not None:
            return anchor_para.insert_paragraph_before()
        return doc.add_paragraph()

    first_entry_para: Optional[Paragraph] = None
    for abbr, expansion in abbreviations:
        para = _new_para()
        if first_entry_para is None:
            first_entry_para = para

        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.first_line_indent = Cm(0)
        para.paragraph_format.tab_stops.add_tab_stop(Cm(4.0), WD_TAB_ALIGNMENT.LEFT)

        run = para.add_run(abbr)
        run.bold = True
        para.add_run(f"\t{expansion}" if expansion else "\t")

    # --- 3. TITLE (before the first entry line) ---
    title_anchor = first_entry_para if first_entry_para is not None else anchor_para
    if title_anchor is not None:
        title_para = title_anchor.insert_paragraph_before("LIST OF ABBREVIATIONS")
    else:
        title_para = doc.add_paragraph("LIST OF ABBREVIATIONS")

    try:
        title_para.style = title_style_name
    except Exception:
        pass
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    print(f"[INFO] Inserted LIST OF ABBREVIATIONS ({len(abbreviations)} entries)")

    # Topmost inserted paragraph = anchor for the next prelim page
    return break_para if break_para is not None else title_para
