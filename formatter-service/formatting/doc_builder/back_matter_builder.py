# formatter-service/formatting/builder/back_matter_builder.py

from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, List, Any
import logging

logger = logging.getLogger(__name__)

class BackMatterBuilder:
    """Builds references and appendices."""
    
    def build(self, doc, structured_doc: Dict[str, Any]):
        """Build all back matter sections."""
        logger.info("Building back matter...")
        
        self._build_references(doc, structured_doc.get('references', []))
        self._build_appendices(doc, structured_doc.get('appendices', []))
    
    def _build_references(self, doc, references: List[str]):
        """Build references section."""
        if references:
            ref_heading = doc.add_paragraph("REFERENCES")
            ref_heading.style = doc.styles['Heading 1']
            ref_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for ref in references:
                ref_para = doc.add_paragraph(ref)
                ref_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            doc.add_page_break()
    
    def _build_appendices(self, doc, appendices: List[Dict[str, Any]]):
        """Build appendices section."""
        if appendices:
            app_heading = doc.add_paragraph("APPENDICES")
            app_heading.style = doc.styles['Heading 1']
            app_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for appendix in appendices:
                self._build_appendix(doc, appendix)
    
    def _build_appendix(self, doc, appendix: Dict[str, Any]):
        """Build a single appendix."""
        label = appendix.get('label', 'Appendix A')
        title = appendix.get('title', 'Untitled Appendix')
        
        appendix_heading = doc.add_paragraph(f"{label}\\n{title}")
        appendix_heading.style = doc.styles['Heading 1']
        
        content = appendix.get('content', [])
        for para_text in content:
            para = doc.add_paragraph(para_text)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY