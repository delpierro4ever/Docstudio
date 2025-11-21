# formatter-service/formatting/builder/chapter_builder.py

from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, List, Any
import logging

logger = logging.getLogger(__name__)

class ChapterBuilder:
    """Builds main chapters and content."""
    
    def build(self, doc, structured_doc: Dict[str, Any]):
        """Build all chapters."""
        logger.info("Building main content...")
        
        chapters = structured_doc.get('chapters', [])
        for chapter in chapters:
            self._build_chapter(doc, chapter)
    
    def _build_chapter(self, doc, chapter: Dict[str, Any]):
        """Build a single chapter with content."""
        chapter_num = chapter.get('number', '1')
        chapter_title = chapter.get('title', 'Untitled Chapter')
        
        # Chapter heading - FIXED: Remove escaped \n
        heading_text = f"CHAPTER {chapter_num}\n{chapter_title.upper()}"
        chapter_heading = doc.add_paragraph(heading_text)
        chapter_heading.style = doc.styles['Heading 1']
        chapter_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Build subsections (Heading 2)
        self._build_subsections(doc, chapter.get('subsections', []))
        
        # Chapter content paragraphs
        paragraphs = chapter.get('paragraphs', [])
        for para_text in paragraphs:
            # FIXED: Remove escaped newlines
            clean_text = para_text.replace('\\n', '\n')
            para = doc.add_paragraph(clean_text)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Build chapter tables
        self._build_chapter_tables(doc, chapter.get('tables', []))
        
        # Build chapter figures
        self._build_chapter_figures(doc, chapter.get('figures', []))
        
        doc.add_page_break()
    
    def _build_subsections(self, doc, subsections: List[Dict[str, Any]]):
        """Build subsection headings (Heading 2)."""
        for subsection in subsections:
            subsection_num = subsection.get('number', '')
            subsection_title = subsection.get('title', 'Untitled Subsection')
            
            # Subsection heading (Heading 2 style)
            if subsection_num:
                heading_text = f"{subsection_num} {subsection_title}"
            else:
                heading_text = subsection_title
                
            subsection_heading = doc.add_paragraph(heading_text)
            subsection_heading.style = doc.styles['Heading 2']  # ✅ Now using Heading 2
            subsection_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Subsection content
            paragraphs = subsection.get('paragraphs', [])
            for para_text in paragraphs:
                clean_text = para_text.replace('\\n', '\n')
                para = doc.add_paragraph(clean_text)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    def _build_chapter_tables(self, doc, tables: List[Dict[str, Any]]):
        """Build tables for this chapter."""
        for table_data in tables:
            label = table_data.get('label', 'Table')
            title = table_data.get('title', '')
            content = table_data.get('content', '')
            
            # Table label and title
            if title:
                table_heading = doc.add_paragraph(f"{label}: {title}")
            else:
                table_heading = doc.add_paragraph(label)
            table_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            table_heading.style = doc.styles['Normal']
            
            # TODO: Create actual table from content
            # For now, just add the content as text
            if content:
                table_content = doc.add_paragraph(content)
                table_content.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    def _build_chapter_figures(self, doc, figures: List[Dict[str, Any]]):
        """Build figures for this chapter."""
        for figure_data in figures:
            label = figure_data.get('label', 'Figure')
            title = figure_data.get('title', '')
            description = figure_data.get('description', '')
            
            # Figure label and title
            if title:
                figure_heading = doc.add_paragraph(f"{label}: {title}")
            else:
                figure_heading = doc.add_paragraph(label)
            figure_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            figure_heading.style = doc.styles['Normal']
            
            # Figure description
            if description:
                figure_desc = doc.add_paragraph(description)
                figure_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Add placeholder for image
            placeholder = doc.add_paragraph("[Image placeholder]")
            placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
            placeholder.italic = True