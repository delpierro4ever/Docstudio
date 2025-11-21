# formatter-service/formatting/builder/front_matter_builder.py

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from typing import Dict, List, Any
import logging

logger = logging.getLogger(__name__)

class FrontMatterBuilder:
    """Builds front matter: title page, TOC, lists, abstract, abbreviations."""
    
    def build(self, doc, structured_doc: Dict[str, Any]):
        """Build all front matter sections."""
        logger.info("Building front matter...")
        
        self._build_title_page(doc, structured_doc.get('meta', {}))
        self._build_abstract(doc, structured_doc.get('preliminaries', {}))
        self._build_table_of_contents_placeholder(doc)
        self._build_list_of_tables_placeholder(doc)
        self._build_list_of_figures_placeholder(doc)
        self._build_abbreviations(doc, structured_doc.get('abbreviations', []))
    
    def _build_title_page(self, doc, meta: Dict[str, Any]):
        """Build title page with document metadata - LEAVE BLANK for UB standards."""
        # UB Standard: Title pages are left blank
        logger.info("Creating blank title pages (UB standard)")
        
        # Create 2 blank pages for title page
        for _ in range(2):
            # Add some empty paragraphs to create blank pages
            for __ in range(10):
                doc.add_paragraph()
            doc.add_page_break()
    
    def _build_abstract(self, doc, preliminaries: Dict[str, Any]):
        """Build abstract section."""
        abstract_text = preliminaries.get('abstract', [])
        if abstract_text:
            abstract_heading = doc.add_paragraph("ABSTRACT")
            abstract_heading.style = doc.styles['Heading 1']
            abstract_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for para_text in abstract_text:
                # Fix: Remove escaped newlines and use actual formatting
                clean_text = para_text.replace('\\n', '\n')
                abstract_para = doc.add_paragraph(clean_text)
                abstract_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            doc.add_page_break()
    
    def _build_table_of_contents_placeholder(self, doc):
        """Add Table of Contents placeholder."""
        toc_heading = doc.add_paragraph("TABLE OF CONTENTS")
        toc_heading.style = doc.styles['Heading 1']
        toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        placeholder = doc.add_paragraph(
            "Note: Please update the Table of Contents in Microsoft Word."
        )
        placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_page_break()
    
    def _build_list_of_tables_placeholder(self, doc):
        """Add List of Tables placeholder."""
        lot_heading = doc.add_paragraph("LIST OF TABLES")
        lot_heading.style = doc.styles['Heading 1']
        lot_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER  # ✅ FIXED: Center aligned
        
        # Add placeholder text
        placeholder = doc.add_paragraph("(List of tables will be generated automatically)")
        placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_page_break()
    
    def _build_list_of_figures_placeholder(self, doc):
        """Add List of Figures placeholder."""
        lof_heading = doc.add_paragraph("LIST OF FIGURES")
        lof_heading.style = doc.styles['Heading 1']
        lof_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER  # ✅ FIXED: Center aligned
        
        # Add placeholder text
        placeholder = doc.add_paragraph("(List of figures will be generated automatically)")
        placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_page_break()
    
    def _build_abbreviations(self, doc, abbreviations: List[Dict[str, Any]]):
        """Build abbreviations section."""
        if abbreviations:
            abbrev_heading = doc.add_paragraph("ABBREVIATIONS")
            abbrev_heading.style = doc.styles['Heading 1']
            abbrev_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER  # ✅ FIXED: Center aligned
            
            for abbrev in abbreviations:
                short = abbrev.get('short', '')
                full = abbrev.get('full', '')
                if short and full:
                    abbrev_para = doc.add_paragraph()
                    # Fix: Use proper tab instead of escaped \t
                    abbrev_run_short = abbrev_para.add_run(f"{short}\t")
                    abbrev_run_short.bold = True
                    abbrev_para.add_run(full)
                    # Align the entire paragraph
                    abbrev_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            doc.add_page_break()