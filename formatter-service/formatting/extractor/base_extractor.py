# formatter-service/formatting/extractor/base_extractor.py

from docx import Document
from typing import Dict, List, Any
from .paragraph_extractor import ParagraphExtractor
from .table_extractor import TableExtractor
from .image_extractor import ImageExtractor

class DocumentExtractor:
    """Main orchestrator for document extraction."""
    
    def __init__(self):
        self.para_extractor = ParagraphExtractor()
        self.table_extractor = TableExtractor()
        self.image_extractor = ImageExtractor()
    
    def extract_for_model(self, input_path: str) -> Dict[str, Any]:
        """
        Extract raw content from DOCX for LLM processing.
        """
        doc = Document(input_path)
        
        return {
            "paragraphs": self.para_extractor.extract(doc),
            "tables": self.table_extractor.extract_for_llm(doc),
            "images": self.image_extractor.extract_for_llm(doc)
        }
    
    def extract_original_tables(self, input_path: str) -> List[Dict[str, Any]]:
        """
        Extract actual table objects with positioning information.
        """
        doc = Document(input_path)
        return self.table_extractor.extract_preserved(doc)
    
    def extract_original_images(self, input_path: str) -> List[Dict[str, Any]]:
        """
        Extract image information with positioning.
        """
        doc = Document(input_path)
        return self.image_extractor.extract_preserved(doc)