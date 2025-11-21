# formatter-service/formatting/extractor/image_extractor.py

import re
from typing import Dict, List, Any
from docx import Document

class ImageExtractor:
    """Handles image extraction for both LLM and preservation."""
    
    def extract_for_llm(self, doc: Document) -> List[Dict[str, Any]]:
        """Detect images and their approximate locations for LLM."""
        images = []
        image_index = 0
        
        # Method 1: Embedded images
        for para_idx, paragraph in enumerate(doc.paragraphs):
            for run in paragraph.runs:
                if self._has_image(run):
                    images.append({
                        "index": image_index,
                        "paragraph_index": para_idx,
                        "description": f"Image {image_index + 1}",
                        "has_caption": self._check_for_image_caption(doc, para_idx),
                        "position_hint": f"image_{image_index}",
                        "type": "embedded"
                    })
                    image_index += 1
        
        # Method 2: Image references in text
        image_refs = self._find_image_references_in_text(doc)
        for ref in image_refs:
            images.append({
                "index": image_index,
                "paragraph_index": ref["paragraph_index"],
                "description": f"Image {image_index + 1}",
                "has_caption": self._check_for_image_caption(doc, ref["paragraph_index"]),
                "position_hint": f"image_{image_index}",
                "reference_text": ref["text"],
                "type": "text_reference"
            })
            image_index += 1
        
        return images
    
    def extract_preserved(self, doc: Document) -> List[Dict[str, Any]]:
        """Extract image information with positioning."""
        preserved_images = []
        image_index = 0
        
        # Method 1: Check for embedded images in runs
        for para_idx, paragraph in enumerate(doc.paragraphs):
            for run in paragraph.runs:
                if self._has_image(run):
                    image_data = {
                        "original_index": image_index,
                        "paragraph_index": para_idx,
                        "run": run,  # Keep the run reference
                        "position": f"paragraph_{para_idx}",
                        "description": f"Image {image_index + 1}",
                        "has_caption": self._check_for_image_caption(doc, para_idx),
                        "approx_size": self._estimate_image_size(run),
                        "detection_method": "embedded"
                    }
                    preserved_images.append(image_data)
                    image_index += 1
        
        # Method 2: Check for image references in text and match with related parts
        image_refs = self._find_image_references_in_text(doc)
        related_images = self._get_related_images(doc)
        
        # Try to match image references with related images
        for ref_idx, image_ref in enumerate(image_refs):
            # Find the best matching related image
            related_image = None
            if ref_idx < len(related_images):
                related_image = related_images[ref_idx]
            
            image_data = {
                "original_index": image_index,
                "paragraph_index": image_ref["paragraph_index"],
                "position": f"paragraph_{image_ref['paragraph_index']}",
                "description": f"Image {image_index + 1}",
                "reference_text": image_ref["text"],
                "has_caption": self._check_for_image_caption(doc, image_ref["paragraph_index"]),
                "detection_method": "text_reference"
            }
            
            if related_image:
                image_data.update({
                    "rel_id": related_image["rel_id"],
                    "rel_part": related_image["rel_part"],
                    "content_type": related_image["content_type"],
                    "detection_method": "text_reference_with_related"
                })
            
            preserved_images.append(image_data)
            image_index += 1
        
        # Method 3: Add any remaining related images that weren't matched
        for i in range(len(image_refs), len(related_images)):
            related_image = related_images[i]
            # Try to find a reasonable position for unmatched related images
            position_para = self._find_best_position_for_image(doc, i)
            
            image_data = {
                "original_index": image_index,
                "paragraph_index": position_para,
                "position": f"paragraph_{position_para}",
                "description": f"Image {image_index + 1}",
                "rel_id": related_image["rel_id"],
                "rel_part": related_image["rel_part"],
                "content_type": related_image["content_type"],
                "has_caption": self._check_for_image_caption(doc, position_para),
                "detection_method": "unmatched_related"
            }
            preserved_images.append(image_data)
            image_index += 1
        
        return preserved_images
    
    def _has_image(self, run) -> bool:
        """Check if a run contains an image."""
        # Method 1: Check for graphic elements
        if run._element.xpath('.//a:graphic'):
            return True
        
        # Method 2: Check for drawing elements
        if run._element.xpath('.//wp:inline') or run._element.xpath('.//wp:anchor'):
            return True
        
        # Method 3: Check for blip elements (image references)
        if run._element.xpath('.//a:blip'):
            return True
        
        return False
    
    def _estimate_image_size(self, run) -> Dict[str, float]:
        """Estimate image size from drawing properties."""
        try:
            # Try to get extent elements for size
            extents = run._element.xpath('.//wp:extent')
            if extents:
                extent = extents[0]
                cx = int(extent.get('cx', 0)) / 914400  # Convert EMU to inches
                cy = int(extent.get('cy', 0)) / 914400
                return {"width_inches": cx, "height_inches": cy}
        except:
            pass
        
        return {"width_inches": 4.0, "height_inches": 3.0}  # Default size
    
    def _find_image_references_in_text(self, doc: Document) -> List[Dict[str, Any]]:
        """Find image references in paragraph text (like ![](media/image1.jpeg))."""
        image_refs = []
        image_patterns = [
            r'!\[.*?\]\(.*?\)',  # Markdown style: ![](path)
            r'Figure\s+\d+:',    # Figure 1:
            r'Image\s+\d+',      # Image 1
            r'media/image\d+\.', # media/image1.jpeg
        ]
        
        for para_idx, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text
            for pattern in image_patterns:
                if re.search(pattern, text, re.IGNORECASE):
                    image_refs.append({
                        "paragraph_index": para_idx,
                        "text": text.strip(),
                        "matched_pattern": pattern
                    })
                    break  # Only count once per paragraph
        
        return image_refs
    
    def _get_related_images(self, doc: Document) -> List[Dict[str, Any]]:
        """Get images from related parts."""
        related_images = []
        
        if hasattr(doc, 'part') and hasattr(doc.part, 'related_parts'):
            for rel_id, rel_part in doc.part.related_parts.items():
                if 'image' in rel_part.content_type:
                    related_images.append({
                        "rel_id": rel_id,
                        "rel_part": rel_part,
                        "content_type": rel_part.content_type
                    })
        
        return related_images
    
    def _find_best_position_for_image(self, doc: Document, image_index: int) -> int:
        """Find a reasonable paragraph position for an unmatched image."""
        # Look for paragraphs that mention figures/images
        for para_idx, paragraph in enumerate(doc.paragraphs):
            text_lower = paragraph.text.lower()
            if any(keyword in text_lower for keyword in ['figure', 'image', 'photo', 'chart']):
                return para_idx
        
        # Fallback: place after table of contents or around middle
        return min(10, len(doc.paragraphs) - 1)
    
    def _check_for_image_caption(self, doc: Document, image_para_idx: int) -> bool:
        """Check if image has a caption in following paragraphs."""
        caption_indicators = ['figure', 'image', 'photo', 'table', 'chart', 'graph']
        
        for i in range(1, 4):  # Check next 3 paragraphs
            if image_para_idx + i < len(doc.paragraphs):
                next_para = doc.paragraphs[image_para_idx + i].text.lower()
                if any(indicator in next_para for indicator in caption_indicators):
                    return True
                    
        return False