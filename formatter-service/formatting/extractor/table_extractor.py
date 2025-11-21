# formatter-service/formatting/extractor/table_extractor.py

from typing import Dict, List, Any
from docx import Document

class TableExtractor:
    """Handles table extraction for both LLM and preservation."""
    
    def extract_for_llm(self, doc: Document) -> List[Dict[str, Any]]:
        """Extract text content from tables for LLM."""
        tables = []
        for table_idx, table in enumerate(doc.tables):
            table_text = []
            
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    table_text.append(" | ".join(row_text))
            
            if table_text:
                tables.append({
                    "index": table_idx,
                    "text": "\n".join(table_text),
                    "row_count": len(table.rows),
                    "column_count": len(table.rows[0].cells) if table.rows else 0,
                    "position_hint": f"table_{table_idx}"
                })
        
        return tables
    
    def extract_preserved(self, doc: Document) -> List[Dict[str, Any]]:
        """Extract actual table objects with positioning information."""
        preserved_tables = []
        
        table_index = 0
        paragraph_count = 0
        
        # Track tables by their position in the document
        for i, element in enumerate(doc.element.body):
            if element.tag.endswith('p'):  # Paragraph
                paragraph_count += 1
            elif element.tag.endswith('tbl'):  # Table
                # Find the corresponding table in python-docx tables list
                if table_index < len(doc.tables):
                    table_obj = doc.tables[table_index]
                    table_data = {
                        "original_index": table_index,
                        "table_object": table_obj,  # Keep the actual table object
                        "position_after_paragraph": paragraph_count,
                        "approx_position": f"after_paragraph_{paragraph_count}",
                        "row_count": len(table_obj.rows),
                        "col_count": len(table_obj.rows[0].cells) if table_obj.rows else 0
                    }
                    preserved_tables.append(table_data)
                    table_index += 1
        
        return preserved_tables