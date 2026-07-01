from __future__ import annotations

from copy import deepcopy
from io import BytesIO
from typing import Any, Dict, List, Optional

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from config.profiles import load_profile
from llm.proofreader import (
    Proofreader,
    collect_proofread_candidates,
    proofreading_enabled,
)
from .abbreviations_builder import extract_abbreviations, insert_list_of_abbreviations
from .basic_style import apply_basic_style
from .field_updater import (
    bake_fields_enabled,
    bake_fields_with_libreoffice,
    enable_update_fields_on_open,
)
from .heading_builder import apply_headings
from .references_builder import format_references_section
from .section_builder import apply_sections
from .table_figure_builder import apply_tables_and_figures, insert_list_of_entries
from .toc_builder import insert_table_of_contents

# Prelim pages we can generate, in their required document order
_GENERATED_PRELIM_ITEMS = [
    "tableOfContents",
    "listOfTables",
    "listOfFigures",
    "abbreviations",
]


def _get_paragraph_index_for_block(
    blocks: List[Dict[str, Any]],
    target_block_id: str,
    doc_paragraphs: List[Paragraph],
) -> Optional[int]:
    """
    Walk through blocks and map paragraph-type blocks to indices in
    doc.paragraphs (same order as the original DOCX).

    ONLY valid while no paragraphs have been inserted/removed — which is
    why format_docx resolves every anchor it needs up front and holds on
    to the Paragraph objects.
    """
    para_index = -1

    for block in blocks:
        if block.get("type") != "paragraph":
            continue

        para_index += 1
        if block.get("id") == target_block_id:
            if 0 <= para_index < len(doc_paragraphs):
                return para_index
            return None

    return None


def _get_paragraph_by_block_id(
    doc: Document,
    blocks: List[Dict[str, Any]],
    block_id: str,
) -> Optional[Paragraph]:
    """Helper to get a Paragraph object by its block ID."""
    para_index = _get_paragraph_index_for_block(blocks, block_id, doc.paragraphs)
    if para_index is not None:
        return doc.paragraphs[para_index]
    return None


def _resolve_main_anchor(
    doc: Document,
    blocks: List[Dict[str, Any]],
    metadata: Dict[str, Any],
) -> Optional[Paragraph]:
    """
    Find the first paragraph of the main content (usually the CHAPTER ONE
    heading). Resolution order:

      1. sections.prelim_ends_before_block_id from the LLM
      2. the first block classified as a chapter_heading
      3. the first paragraph block classified into section "main"
    """
    boundary_id = metadata.get("sections", {}).get("prelim_ends_before_block_id")
    if boundary_id:
        para = _get_paragraph_by_block_id(doc, blocks, boundary_id)
        if para is not None:
            return para

    block_meta: Dict[str, Any] = metadata.get("blocks", {}) or {}

    for key in ("chapter_heading", None):
        for block in blocks:
            if block.get("type") != "paragraph":
                continue
            meta = block_meta.get(block.get("id"), {}) or {}
            if key == "chapter_heading" and meta.get("role") == "chapter_heading":
                return _get_paragraph_by_block_id(doc, blocks, block["id"])
            if key is None and meta.get("section") == "main":
                return _get_paragraph_by_block_id(doc, blocks, block["id"])

    return None


def _apply_proofreading(
    doc: Document,
    blocks: List[Dict[str, Any]],
    metadata: Dict[str, Any],
) -> None:
    """
    Run the LLM grammar/spelling pass over body paragraphs and write the
    corrected text back. Character-level run formatting inside a corrected
    paragraph is collapsed to the first run's formatting (acceptable for
    body text, whose look is governed by the global profile anyway).

    Best-effort: any failure leaves the document untouched.
    """
    candidates = collect_proofread_candidates(blocks, metadata)
    if not candidates:
        return

    corrections = Proofreader().correct_texts(candidates)
    if not corrections:
        print("[INFO] Proofreading: no corrections applied")
        return

    applied = 0
    for block_id, corrected in corrections.items():
        paragraph = _get_paragraph_by_block_id(doc, blocks, block_id)
        if paragraph is None:
            continue

        rPr = None
        if paragraph.runs:
            existing_rPr = paragraph.runs[0]._r.find(qn("w:rPr"))
            if existing_rPr is not None:
                rPr = deepcopy(existing_rPr)

        paragraph.clear()
        run = paragraph.add_run(corrected)
        if rPr is not None:
            run._r.insert(0, rPr)

        # Keep the block list in sync so later passes (e.g. abbreviation
        # extraction) see the corrected text.
        for block in blocks:
            if block.get("id") == block_id:
                block["text"] = corrected
                break

        applied += 1

    print(f"[INFO] Proofreading: corrected {applied} paragraph(s)")


def format_docx(
    input_path: str,
    blocks: List[Dict[str, Any]],
    metadata: Dict[str, Any],
    profile_id: Optional[str] = None,
) -> bytes:
    """
    Main entry point for building the final formatted DOCX.

    Ordering matters: every pass that maps block ids onto doc.paragraphs
    (steps 1-4 and anchor resolution) runs BEFORE any paragraph is
    inserted, because insertion shifts the mapping.
    """

    profile = load_profile(profile_id)
    doc = Document(input_path)

    # 1. Global style + heading styles (indices aligned)
    try:
        apply_basic_style(doc, profile, metadata)
    except Exception as exc:
        print(f"[WARN] apply_basic_style failed: {exc}")

    try:
        apply_headings(doc, blocks, metadata, profile)
    except Exception as exc:
        print(f"[WARN] apply_headings failed: {exc}")

    # 2. Grammar/spelling corrections (text-only; indices stay aligned)
    grammar_enabled = (profile.get("grammar", {}) or {}).get("enabled", True)
    if grammar_enabled and proofreading_enabled():
        try:
            _apply_proofreading(doc, blocks, metadata)
        except Exception as exc:
            print(f"[WARN] proofreading failed: {exc}")

    # 3. Normalize captions with SEQ fields (indices aligned)
    try:
        apply_tables_and_figures(doc, blocks, metadata, profile)
    except Exception as exc:
        print(f"[WARN] apply_tables_and_figures failed: {exc}")

    # 4. References formatting (indices aligned / text scan)
    try:
        format_references_section(doc, blocks, metadata, profile)
    except Exception as exc:
        print(f"[WARN] format_references_section raised unexpectedly: {exc}")

    # 5. Resolve the main-content anchor ONCE, while indices are aligned.
    #    The Paragraph object stays valid through the insertions below
    #    (they all land before it).
    main_anchor = _resolve_main_anchor(doc, blocks, metadata)
    if main_anchor is None:
        print("[WARN] No main-content boundary found; prelim pages will be "
              "inserted at the top and page numbering will be skipped")

    # 6. Collect abbreviations BEFORE inserting anything (uses block text)
    abbreviations = []
    abbrev_enabled = (profile.get("abbreviations", {}) or {}).get("autoGenerate", True)
    if abbrev_enabled:
        try:
            abbreviations = extract_abbreviations(blocks, metadata)
        except Exception as exc:
            print(f"[WARN] abbreviation extraction failed: {exc}")

    # 7. Insert generated prelim pages. Required order in the document is
    #    TOC -> LOT -> LOF -> Abbreviations, so we insert in REVERSE order,
    #    each time just before the current anchor.
    fallback_anchor = doc.paragraphs[0] if doc.paragraphs else None
    current_anchor = main_anchor if main_anchor is not None else fallback_anchor

    order = (profile.get("structure", {}) or {}).get("order") or []
    generated_items = [k for k in order if k in _GENERATED_PRELIM_ITEMS]
    if not generated_items:
        generated_items = list(_GENERATED_PRELIM_ITEMS)

    for item_key in reversed(generated_items):
        try:
            inserted_title_para = None

            if item_key == "tableOfContents":
                inserted_title_para = insert_table_of_contents(
                    doc, metadata, anchor_para=current_anchor,
                )
            elif item_key == "listOfTables":
                inserted_title_para = insert_list_of_entries(
                    doc, metadata,
                    title="LIST OF TABLES",
                    caption_label="Table",
                    anchor_para=current_anchor,
                )
            elif item_key == "listOfFigures":
                inserted_title_para = insert_list_of_entries(
                    doc, metadata,
                    title="LIST OF FIGURES",
                    caption_label="Figure",
                    anchor_para=current_anchor,
                )
            elif item_key == "abbreviations":
                inserted_title_para = insert_list_of_abbreviations(
                    doc, abbreviations, anchor_para=current_anchor,
                )

            # The newly inserted title becomes the anchor for the next
            # (earlier) prelim page.
            if inserted_title_para is not None:
                current_anchor = inserted_title_para

        except Exception as exc:
            print(f"[WARN] Failed to insert {item_key}: {exc}")

    # 8. Section break + page numbering (roman prelim / arabic main).
    #    Runs last so the break lands after every generated prelim page.
    try:
        apply_sections(doc, main_anchor)
    except Exception as exc:
        print(f"[WARN] apply_sections failed: {exc}")

    # 9. Make Word refresh all fields (TOC/LOT/LOF/SEQ/PAGE) on open.
    try:
        enable_update_fields_on_open(doc)
    except Exception as exc:
        print(f"[WARN] enable_update_fields_on_open failed: {exc}")

    # 10. Save; optionally bake field results via headless LibreOffice so
    #     the TOC is populated even in viewers that ignore updateFields.
    bio = BytesIO()
    doc.save(bio)
    output = bio.getvalue()

    if bake_fields_enabled():
        baked = bake_fields_with_libreoffice(output)
        if baked:
            print("[INFO] Field results baked via LibreOffice")
            output = baked
        else:
            print("[WARN] Field bake unavailable; relying on update-on-open")

    return output
