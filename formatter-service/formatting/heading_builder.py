from typing import Dict, Any, List, Optional

from docx import Document
from docx.text.paragraph import Paragraph
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def apply_headings(
    doc: Document,
    blocks: List[Dict[str, Any]],
    metadata: Dict[str, Any],
    profile: Dict[str, Any],
) -> None:
    """
    Apply heading styles (Heading 1, 2, 3) to paragraphs
    based on LLM metadata + profile.heading config.
    """
    meta_blocks: Dict[str, Any] = metadata.get("blocks", {})
    headings_cfg: Dict[str, Any] = profile.get("headings", {})

    # Convenience handles for each level
    chapter_cfg = headings_cfg.get("chapterHeading", {})
    sub1_cfg = headings_cfg.get("subHeading1", {})
    sub2_cfg = headings_cfg.get("subHeading2", {})

    # Map paragraph-type blocks to doc.paragraphs indices
    para_index = -1

    for block in blocks:
        if block.get("type") != "paragraph":
            continue

        para_index += 1
        if para_index >= len(doc.paragraphs):
            break

        block_id = block.get("id")
        meta = meta_blocks.get(block_id, {})
        role = meta.get("role", "")

        paragraph: Paragraph = doc.paragraphs[para_index]

        if role == "chapter_heading":
            _normalize_heading_text(paragraph, block, meta, profile)
            _apply_heading_style(paragraph, chapter_cfg, profile) # <-- PASSING PROFILE
        elif role == "section_heading":
            _apply_heading_style(paragraph, sub1_cfg, profile) # <-- PASSING PROFILE
        elif role == "subsection_heading":
            _apply_heading_style(paragraph, sub2_cfg, profile) # <-- PASSING PROFILE
        # else: leave as body paragraph

def _normalize_heading_text(
    paragraph: Paragraph,
    block: Dict[str, Any],
    meta: Dict[str, Any],
    profile: Dict[str, Any],
) -> None:
    """
    Enforce consistent text formatting for chapter headings (e.g., CHAPTER X: Title).
    """
    
    chapter_num = meta.get("chapter")
    raw_title = meta.get("title", paragraph.text.strip())

    if not chapter_num:
        return

    chapter_prefix = profile.get("headings", {}).get("chapterHeading", {}).get("detectPattern", "CHAPTER")
    
    if raw_title and chapter_num is not None:
        
        # Basic cleanup of the raw title text (removes existing numbering/prefix)
        cleaned_title = raw_title
        title_parts = raw_title.split(':')
        if len(title_parts) > 1:
            cleaned_title = title_parts[1].strip()
        elif ' ' in raw_title:
            try:
                # Try removing the first two words (e.g., 'Chapter 1')
                cleaned_title = ' '.join(raw_title.split(' ')[2:]).strip() 
            except IndexError:
                pass

        # If cleanup resulted in an empty string, fallback to raw text
        if not cleaned_title:
            cleaned_title = raw_title.strip()
        
        # Construct the normalized text
        normalized_text = f"{chapter_prefix.upper()} {chapter_num}: {cleaned_title}"
        
        # Apply the normalized text to the paragraph
        # This overwrites the run text, ensuring consistency
        paragraph.text = normalized_text


def _apply_heading_style(paragraph: Paragraph, cfg: Dict[str, Any], profile: Dict[str, Any]) -> None:
    """
    Apply a single heading style configuration to a paragraph.
    cfg comes from profile["headings"][...].
    """
    if not cfg:
        return

    # --- Style name (e.g. "Heading 1") ---
    style_name = cfg.get("styleName")
    if style_name:
        # We rely on basic_style.py to guarantee this style exists
        try:
            paragraph.style = style_name
        except KeyError:
            pass # Should not fail if basic_style.py ran correctly

    # --- Alignment ---
    alignment = cfg.get("alignment", "left").lower()
    if alignment == "center":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == "right":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    fmt = paragraph.paragraph_format

    # --- Spacing ---
    spacing_before = cfg.get("spacingBefore")
    spacing_after = cfg.get("spacingAfter")

    if spacing_before is not None:
        fmt.space_before = Pt(spacing_before)
    if spacing_after is not None:
        fmt.space_after = Pt(spacing_after)

    # --- Font-level changes on runs ---
    font_size = cfg.get("fontSize")  # in points
    bold = cfg.get("bold", False)
    italics = cfg.get("italics", False)
    all_caps = cfg.get("allCaps", False)
    
    global_font_family = profile.get("font", {}).get("family") 

    for run in paragraph.runs:
        if font_size is not None:
            run.font.size = Pt(font_size)

        if global_font_family:
             run.font.name = global_font_family
             
        run.font.bold = bold
        run.font.italic = italics

        if all_caps:
            run.text = run.text.upper()