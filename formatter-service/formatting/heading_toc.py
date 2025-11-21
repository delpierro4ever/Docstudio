# formatter-service/formatting/heading_toc.py

import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Pt
from docx.oxml.ns import qn

# CHAPTER 1, CHAPTER 2: ...
CHAPTER_HEADING_REGEX = re.compile(
    r"^\s*chapter\s+\d+[:.\-\s]*",
    re.IGNORECASE,
)

# 1.1 Introduction, 2.3 Background, etc.
SUBHEADING_L2_REGEX = re.compile(
    r"^\s*\d+\.\d+\s+.+",
    re.IGNORECASE,
)

# 1.1.1 Something, 2.3.4 More details, etc.
SUBHEADING_L3_REGEX = re.compile(
    r"^\s*\d+\.\d+\.\d+\s+.+",
    re.IGNORECASE,
)

HEADING_FONT_NAME = "Times New Roman"


def _force_run_font(run, size_pt: int, bold: bool = True, all_caps: bool = False, italic: bool = False):
    """Force consistent font on a run."""
    font = run.font
    font.name = HEADING_FONT_NAME
    font.size = Pt(size_pt)
    font.bold = bold
    font.italic = italic
    try:
        font.all_caps = all_caps
    except Exception:
        pass

    if font.element.rPr is not None:
        rPr = font.element.rPr
        rFonts = rPr.rFonts
        rFonts.set(qn("w:ascii"), HEADING_FONT_NAME)
        rFonts.set(qn("w:hAnsi"), HEADING_FONT_NAME)
        rFonts.set(qn("w:cs"), HEADING_FONT_NAME)


def apply_heading_if_needed(paragraph):
    """
    Detect 'Chapter 1 ...' etc. and style as a main chapter heading:
    - Heading 1 (if exists)
    - Centered
    - Bold
    - ALL CAPS
    - 14pt
    """
    text = paragraph.text.strip()
    if not text:
        return

    if not CHAPTER_HEADING_REGEX.match(text):
        return

    try:
        # Try apply Heading 1 style if it exists
        try:
            paragraph.style = paragraph.part.document.styles["Heading 1"]
        except Exception:
            pass

        pf = paragraph.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.space_before = Pt(24)
        pf.space_after = Pt(12)

        for run in paragraph.runs:
            _force_run_font(run, size_pt=14, bold=True, all_caps=True, italic=False)

    except Exception as e:
        print("Error in apply_heading_if_needed:", e)


def apply_subheading_if_needed(paragraph):
    """
    Detect numeric subheadings:
      - Heading 2: '1.1 Title'
      - Heading 3: '1.1.1 Title'
    and style them.
    """
    text = paragraph.text.strip()
    if not text:
        return

    try:
        # First check for level 3 (1.1.1 ...)
        if SUBHEADING_L3_REGEX.match(text):
            # Try Heading 3
            try:
                paragraph.style = paragraph.part.document.styles["Heading 3"]
            except Exception:
                pass

            pf = paragraph.paragraph_format
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.space_before = Pt(6)
            pf.space_after = Pt(3)

            for run in paragraph.runs:
                _force_run_font(run, size_pt=12, bold=False, all_caps=False, italic=True)
            return

        # Then check for level 2 (1.1 ...)
        if SUBHEADING_L2_REGEX.match(text):
            # Try Heading 2
            try:
                paragraph.style = paragraph.part.document.styles["Heading 2"]
            except Exception:
                pass

            pf = paragraph.paragraph_format
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.space_before = Pt(12)
            pf.space_after = Pt(6)

            for run in paragraph.runs:
                _force_run_font(run, size_pt=12, bold=True, all_caps=False, italic=False)
            return

    except Exception as e:
        print("Error in apply_subheading_if_needed:", e)


def insert_toc_title(doc: Document):
    """
    Insert a simple 'TABLE OF CONTENTS' title at the very top.
    User can then manually insert TOC in Word if needed.
    """
    if not doc.paragraphs:
        return

    first_p = doc.paragraphs[0]

    title_p = first_p.insert_paragraph_before("TABLE OF CONTENTS")

    # Try TOC Heading, then Heading 1, else default
    try:
        title_p.style = doc.styles["TOC Heading"]
    except Exception:
        try:
            title_p.style = doc.styles["Heading 1"]
        except Exception:
            pass

    pf = title_p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before = Pt(24)
    pf.space_after = Pt(12)

    for run in title_p.runs:
        _force_run_font(run, size_pt=14, bold=True, all_caps=True, italic=False)


def page_break_before_first_chapter(doc: Document):
    """
    Ensure the first 'Chapter X...' heading starts on a new page,
    by inserting a page-break paragraph immediately before it.
    """
    first_chapter_para = None

    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if CHAPTER_HEADING_REGEX.match(text):
            first_chapter_para = p
            break

    if first_chapter_para is None:
        return

    # Insert a paragraph BEFORE the chapter and add a page break
    break_p = first_chapter_para.insert_paragraph_before()
    break_p.add_run().add_break(WD_BREAK.PAGE)
