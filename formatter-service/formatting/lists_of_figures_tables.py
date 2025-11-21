# formatter-service/formatting/lists_of_figures_tables.py

from typing import List
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK


def collect_tables_and_figures(doc: Document):
    """
    Collect captions AFTER formatting.

    We assume:
      - Table captions start with 'Table N:'
      - Figure captions start with 'Figure N:'

    So we just scan paragraphs for lines starting with 'Table ' or 'Figure '.
    """
    tables: List[str] = []
    figures: List[str] = []

    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue

        lower = text.lower()
        if lower.startswith("table "):
            tables.append(text)
        elif lower.startswith("figure "):
            figures.append(text)

    # Remove duplicates while preserving order
    tables = list(dict.fromkeys(tables))
    figures = list(dict.fromkeys(figures))

    return tables, figures


def _find_toc_paragraph(doc: Document):
    """Find the first 'TABLE OF CONTENTS' paragraph."""
    for p in doc.paragraphs:
        if (p.text or "").strip().upper() == "TABLE OF CONTENTS":
            return p
    return None


def insert_lists_of_tables_and_figures(
    doc: Document,
    tables: List[str],
    figures: List[str],
) -> None:
    """
    Insert after the TOC, each on its own page, in this order:

      [Page] TABLE OF CONTENTS
      [Page] LIST OF TABLES  (if any)
      [Page] LIST OF FIGURES (if any)
    """
    if not tables and not figures:
        return

    toc_p = _find_toc_paragraph(doc)
    if toc_p is None:
        return

    anchor_el = toc_p._p
    current = anchor_el

    # --- Page break after TOC ---
    pb1_para = doc.add_paragraph()
    pb1_run = pb1_para.add_run()
    pb1_run.add_break(WD_BREAK.PAGE)
    pb1_el = pb1_para._p

    current.addnext(pb1_el)
    current = pb1_el

    # --- LIST OF TABLES ---
    if tables:
        lot_title_p = doc.add_paragraph("LIST OF TABLES")
        try:
            lot_title_p.style = doc.styles["Heading 1"]
        except Exception:
            pass

        lot_pf = lot_title_p.paragraph_format
        lot_pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lot_pf.space_before = Pt(24)
        lot_pf.space_after = Pt(12)

        for run in lot_title_p.runs:
            run.font.bold = True
            run.font.size = Pt(14)

        lot_title_el = lot_title_p._p
        current.addnext(lot_title_el)
        current = lot_title_el

        for caption in tables:
            p = doc.add_paragraph(caption)
            pf = p.paragraph_format
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.space_before = Pt(4)
            pf.space_after = Pt(4)

            for run in p.runs:
                run.font.size = Pt(12)

            el = p._p
            current.addnext(el)
            current = el

    # --- LIST OF FIGURES ---
    if figures:
        # Page break before LOF
        pb2_para = doc.add_paragraph()
        pb2_run = pb2_para.add_run()
        pb2_run.add_break(WD_BREAK.PAGE)
        pb2_el = pb2_para._p

        current.addnext(pb2_el)
        current = pb2_el

        lof_title_p = doc.add_paragraph("LIST OF FIGURES")
        try:
            lof_title_p.style = doc.styles["Heading 1"]
        except Exception:
            pass

        lof_pf = lof_title_p.paragraph_format
        lof_pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lof_pf.space_before = Pt(24)
        lof_pf.space_after = Pt(12)

        for run in lof_title_p.runs:
            run.font.bold = True
            run.font.size = Pt(14)

        lof_title_el = lof_title_p._p
        current.addnext(lof_title_el)
        current = lof_title_el

        for caption in figures:
            p = doc.add_paragraph(caption)
            pf = p.paragraph_format
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.space_before = Pt(4)
            pf.space_after = Pt(4)

            for run in p.runs:
                run.font.size = Pt(12)

            el = p._p
            current.addnext(el)
            current = el
