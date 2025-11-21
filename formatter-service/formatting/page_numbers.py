# formatter-service/formatting/page_numbers.py

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def _add_page_field_to_paragraph(paragraph):
    """
    Insert a PAGE field into the given paragraph so Word shows the current page number.
    """
    # Ensure the paragraph has at least one run
    run = paragraph.add_run()

    # <w:fldSimple w:instr="PAGE" />
    fld_simple = OxmlElement("w:fldSimple")
    fld_simple.set(qn("w:instr"), "PAGE")

    # Attach the field to the underlying run element
    run._r.append(fld_simple)


def _ensure_footer_paragraph(section):
    """
    Get a paragraph in the footer to host the page number.
    If the first footer paragraph is empty, reuse it; otherwise create a new one.
    """
    footer = section.footer

    if footer.paragraphs:
        para = footer.paragraphs[0]
        if para.text.strip():
            # First paragraph has text → create a dedicated one for page number
            para = footer.add_paragraph()
    else:
        para = footer.add_paragraph()

    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return para


def add_page_numbers(doc):
    """
    Add centered page numbers to the footer of every section using the PAGE field.
    (For now: same Arabic numbering across the whole document.)
    """
    for section in doc.sections:
        para = _ensure_footer_paragraph(section)
        _add_page_field_to_paragraph(para)
