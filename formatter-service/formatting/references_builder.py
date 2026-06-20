from __future__ import annotations

from typing import List, Dict, Any, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm


def format_references_section(
    doc: Document,
    blocks: List[Dict[str, Any]],
    metadata: Dict[str, Any],
    profile: Dict[str, Any],
) -> None:
    """
    MVP implementation for formatting the REFERENCES section.

    Includes the CRITICAL STEP: Applying hanging indents to all paragraphs
    following the REFERENCES heading until the end of the document.
    """

    try:
        structure_meta = metadata.get("structure", {})
        references_meta = structure_meta.get("references", {})

        title_block_id: Optional[str] = references_meta.get("title_block_id")
        title_text_target: str = _get_references_title_text_from_profile(profile)

        # 1) If we have an explicit title_block_id, map it to a paragraph
        found_para_index = None
        if title_block_id:
            found_para_index = _get_paragraph_index_for_block(blocks, title_block_id, doc.paragraphs)
            if found_para_index is not None:
                _style_references_heading(doc.paragraphs[found_para_index], title_text_target)
                _apply_references_entry_formatting(doc, found_para_index, profile)
                return

        # 2) Fallback: scan all paragraphs for one that matches the title text
        normalized_target = title_text_target.strip().upper()

        for i, para in enumerate(doc.paragraphs):
            if para.text.strip().upper() == normalized_target:
                _style_references_heading(para, title_text_target)
                # Apply entry formatting starting from the paragraph *after* the heading
                _apply_references_entry_formatting(doc, i, profile)
                return

        # If we reach here, we simply didn't find a references heading.
        return

    except Exception as exc:
        # Absolutely do NOT break the whole pipeline because references failed.
        print(f"[WARN] format_references_section failed: {exc}")
        return


def _get_paragraph_index_for_block(
    blocks: List[Dict[str, Any]],
    target_block_id: str,
    doc_paragraphs,
) -> Optional[int]:
    """
    Walk through blocks and map paragraph-type blocks to
    indices in doc.paragraphs (same order as original DOCX).
    """
    from docx.text.paragraph import Paragraph

    para_index = -1

    for block in blocks:
        if block.get("type") != "paragraph":
            continue

        para_index += 1
        if block.get("id") == target_block_id:
            if 0 <= para_index < len(doc_paragraphs):
                return para_index
            return None

    return None


def _get_references_title_text_from_profile(profile: Dict[str, Any]) -> str:
    """
    Extract the expected references title text from the profile JSON.
    Fallback to 'REFERENCES' if not present.
    """
    structure = profile.get("structure", {})
    refs_cfg = structure.get("references", {})
    return refs_cfg.get("titleText", "REFERENCES")


def _style_references_heading(para, title_text: str) -> None:
    """
    Apply a basic heading style to the REFERENCES title paragraph:
      - Set text to title_text (uppercased)
      - Bold all runs
      - Center alignment
      - Extra spacing before/after
    """
    # Set the text
    para.text = title_text.upper()

    # Make all runs bold
    for run in para.runs:
        run.bold = True

    # Center alignment
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Spacing (values are in twips: 20 * points)
    pf = para.paragraph_format
    # Using 12pt before/6pt after is a reasonable default
    pf.space_before = 12 * 20
    pf.space_after = 6 * 20


def _apply_references_entry_formatting(
    doc: Document,
    heading_index: int,
    profile: Dict[str, Any],
) -> None:
    """
    Apply hanging indent formatting to all paragraphs following the
    references heading until the end of the document.
    """
    # Get configuration for hanging indent value
    structure_cfg = profile.get("structure", {})
    refs_cfg = structure_cfg.get("references", {})
    # Use 1.0 cm as the default if config is missing (common standard)
    hanging_indent_cm = refs_cfg.get("hangingIndentCm", 1.0)
    
    # We start *after* the references heading
    start_index = heading_index + 1
    
    # Iterate through all remaining paragraphs in the document
    for i in range(start_index, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        pf = para.paragraph_format
        
        # Apply the hanging indent rule:
        # 1. Left indent sets the position of the second and subsequent lines.
        # 2. Hanging indent moves the first line back to the margin.
        try:
            # Set the entire paragraph block left by hanging_indent_cm
            pf.left_indent = Cm(hanging_indent_cm)
            # Pull the first line back by the same amount
            pf.first_line_indent = Cm(-hanging_indent_cm)
        except Exception as exc:
            print(f"[WARN] Failed to apply hanging indent for paragraph {i}: {exc}")