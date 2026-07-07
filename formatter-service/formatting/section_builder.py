from typing import Optional

from copy import deepcopy

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph


def apply_sections(
    doc: Document,
    main_anchor: Optional[Paragraph],
) -> None:
    """
    Split the document into two sections at the prelim/main boundary and
    apply page numbering:

      - Section 1 (preliminaries): lowercase roman numerals (i, ii, iii)
      - Section 2 (main content):  arabic numerals restarting at 1

    `main_anchor` is the FIRST paragraph of the main content (e.g. the
    CHAPTER ONE heading), resolved as a live Paragraph object by the caller
    BEFORE any prelim pages were inserted.

    The old implementation recomputed the boundary from the original block
    list AFTER the TOC/LOT/LOF insertion had already shifted
    doc.paragraphs, so the break landed on the wrong paragraph. Anchoring
    on the paragraph object itself makes the placement immune to
    insertions (they all happen before the anchor, never on it).
    """
    if main_anchor is None:
        print("[WARN] apply_sections: no main-content anchor; skipping page numbering")
        return

    # 1) Insert a paragraph carrying a <w:sectPr> immediately before the
    #    main content. A paragraph-level sectPr CLOSES a section, so this
    #    defines the preliminary section; the body-level sectPr continues
    #    to describe the main section.
    sect_para = main_anchor.insert_paragraph_before()
    pPr = sect_para._p.get_or_add_pPr()
    sectPr = _build_sectPr_from_body(doc)
    pPr.append(sectPr)

    # 2) Apply numbering formats + PAGE fields per section.
    sections = list(doc.sections)
    if len(sections) < 2:
        print("[WARN] apply_sections: expected at least 2 sections after split")
        return

    _apply_prelim_page_numbering(sections[0])

    first_main = True
    for sec in sections[1:]:
        _apply_main_page_numbering(sec, restart_at_1=first_main)
        first_main = False

    print("[INFO] Section break inserted at prelim/main boundary "
          "(roman prelim pages, arabic main pages)")


def _build_sectPr_from_body(doc: Document) -> OxmlElement:
    """
    Build a <w:sectPr> for the preliminary section by cloning the page
    geometry (size, margins, columns) of the document's body sectPr, so
    prelim pages keep the same layout instead of reverting to defaults.
    """
    sectPr = OxmlElement("w:sectPr")

    body_sectPr = doc.element.body.find(qn("w:sectPr"))
    if body_sectPr is not None:
        for tag in ("w:pgSz", "w:pgMar", "w:cols", "w:docGrid"):
            el = body_sectPr.find(qn(tag))
            if el is not None:
                sectPr.append(deepcopy(el))

    type_el = OxmlElement("w:type")
    type_el.set(qn("w:val"), "nextPage")
    sectPr.insert(0, type_el)

    return sectPr


# ----------------------------------------------------------------------
# Page numbering + footer helpers
# ----------------------------------------------------------------------


def _apply_prelim_page_numbering(section) -> None:
    """Roman numerals (i, ii, iii) in the preliminary section."""
    footer = section.footer

    try:
        footer.is_linked_to_previous = False
    except AttributeError:
        pass

    _set_page_number_format(section, fmt="lowerRoman", start=1)

    _ensure_page_field_in_footer(footer)


def _apply_main_page_numbering(section, restart_at_1: bool = False) -> None:
    """Arabic numbering in main sections."""
    footer = section.footer

    try:
        footer.is_linked_to_previous = False
    except AttributeError:
        pass

    start_val: Optional[int] = 1 if restart_at_1 else None
    _set_page_number_format(section, fmt="decimal", start=start_val)

    _ensure_page_field_in_footer(footer)


def _set_page_number_format(section, fmt: str, start: Optional[int] = None) -> None:
    """Centralized way to set <w:pgNumType> for a section."""
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn("w:pgNumType"))
    if pgNumType is None:
        pgNumType = OxmlElement("w:pgNumType")
        sectPr.append(pgNumType)

    pgNumType.set(qn("w:fmt"), fmt)

    if start is not None:
        pgNumType.set(qn("w:start"), str(start))
    else:
        if pgNumType.get(qn("w:start")) is not None:
            del pgNumType.attrib[qn("w:start")]


def add_arabic_page_numbers(doc: Document) -> None:
    """
    Add simple Arabic (decimal) page numbering to every section in the
    document — used by the quick/print-ready pipeline which has no
    prelim/main split.
    """
    for i, section in enumerate(doc.sections):
        _apply_main_page_numbering(section, restart_at_1=(i == 0))


def _ensure_page_field_in_footer(footer) -> None:
    """
    Insert a { PAGE } field in the footer, right-aligned, if not already
    present.
    """
    para = None
    for existing in footer.paragraphs:
        if any(
            fld is not None
            for fld in existing._p.iter(qn("w:fldChar"))
        ):
            return  # a field already lives here
        para = existing

    if para is None:
        para = footer.add_paragraph()

    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Build complex field: { PAGE }
    begin_run = para.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(fld_char_begin)

    instr_run = para.add_run()
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    instr_run._r.append(instr_text)

    end_run = para.add_run()
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    end_run._r.append(fld_char_end)
