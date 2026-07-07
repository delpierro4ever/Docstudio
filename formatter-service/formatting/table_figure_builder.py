from typing import Dict, Any, List, Optional

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


# ---------------------------------------------------------------------------
# LOT / LOF insertion (generated prelim pages)
# ---------------------------------------------------------------------------

def _add_list_field(paragraph: Paragraph, caption_label: str) -> None:
    """
    Inserts a functional List field (LOT/LOF).

    The field instruction `TOC \\c "<label>"` collects every caption built
    with a `SEQ <label>` field (see write_caption_with_seq below) — which is
    why captions must contain real SEQ fields for these lists to populate.
    """
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 1) Begin field
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    fld_char_begin.set(qn("w:dirty"), "true")
    run._r.append(fld_char_begin)

    # 2) Instruction text
    run = paragraph.add_run()
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = f' TOC \\h \\z \\c "{caption_label}" '
    run._r.append(instr_text)

    # 3) Separate
    run = paragraph.add_run()
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_char_sep)

    # 4) Placeholder result run (replaced when fields are updated)
    placeholder_run = paragraph.add_run()
    placeholder_run.text = "Right-click and choose 'Update Field' if this text is still visible."
    placeholder_run.italic = True

    # 5) End
    run = paragraph.add_run()
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_end)


def insert_list_of_entries(
    doc: Document,
    metadata: Dict[str, Any],
    title: str,
    caption_label: str,
    anchor_para: Optional[Paragraph] = None,
) -> Optional[Paragraph]:
    """
    Inserts a List of Tables/Figures title and field BEFORE the anchor_para,
    preceded by a simple page break.

    Section breaks are handled once by section_builder at the prelim/main
    boundary. (The old add_break(WD_SECTION.NEXT_PAGE) call was invalid and
    crashed the List of Figures on every run.)

    Returns the topmost inserted paragraph (the new anchor for the next
    prelim page inserted above this one).
    """
    title_style_name = metadata.get("toc", {}).get("titleStyleName", "Heading 1")

    # --- 1. PAGE BREAK ---
    is_first_paragraph_in_doc = (
        anchor_para is not None
        and doc.paragraphs
        and anchor_para._p is doc.paragraphs[0]._p
    )

    break_para = None
    if anchor_para is not None and not is_first_paragraph_in_doc:
        break_para = anchor_para.insert_paragraph_before()
        break_para.add_run().add_break(WD_BREAK.PAGE)

    # --- 2. INSERT LIST FIELD (just before the anchor) ---
    if anchor_para is not None:
        list_para = anchor_para.insert_paragraph_before()
    else:
        list_para = doc.add_paragraph()

    _add_list_field(list_para, caption_label)

    # --- 3. INSERT TITLE (just before the field) ---
    title_para = list_para.insert_paragraph_before(title)

    try:
        title_para.style = title_style_name
    except Exception:
        pass
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    print(f"[INFO] Inserted {title} field")

    # Topmost inserted paragraph = anchor for the next prelim page
    return break_para if break_para is not None else title_para


# ---------------------------------------------------------------------------
# Caption normalization with SEQ fields
# ---------------------------------------------------------------------------

def _get_paragraph_by_block_id(
    doc: Document,
    blocks: List[Dict[str, Any]],
    block_id: str,
) -> Optional[Paragraph]:
    """
    Map a P# block id to the corresponding paragraph by walking the block
    list in order (paragraph-type blocks map 1:1 onto doc.paragraphs).

    The old implementation did doc.paragraphs[int(block_id[1:])], which was
    off by one (P# is 1-based) and ignored the block/paragraph mapping
    entirely, so captions were never found.
    """
    para_index = -1
    for block in blocks:
        if block.get("type") != "paragraph":
            continue
        para_index += 1
        if block.get("id") == block_id:
            if 0 <= para_index < len(doc.paragraphs):
                return doc.paragraphs[para_index]
            return None
    return None


def _ensure_caption_style(doc: Document) -> None:
    """Make sure the built-in 'Caption' style exists in this document."""
    try:
        doc.styles["Caption"]
    except KeyError:
        style = doc.styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
        try:
            style.base_style = doc.styles["Normal"]
        except KeyError:
            pass


def _append_seq_field(paragraph: Paragraph, label: str, cached_number: str) -> None:
    """
    Append a `{ SEQ <label> \\* ARABIC \\s 1 }` field to the paragraph.

    The `\\s 1` switch restarts numbering after each Heading 1 (chapter),
    which yields the per-chapter index used in "Table 2.3". The cached
    result we write is the number we computed ourselves, so the caption
    looks right even before fields are refreshed.
    """
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" SEQ {label} \\* ARABIC \\s 1 "
    run._r.append(instr)

    run = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    paragraph.add_run(cached_number)

    run = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def write_caption_with_seq(
    paragraph: Paragraph,
    label: str,
    chapter: Optional[int],
    index_in_chapter: int,
    caption_text: str,
    profile: Dict[str, Any],
    role: str,
) -> None:
    """
    Rewrite a caption paragraph as e.g.:

        Table 2.{SEQ Table \\s 1 -> 3}: Distribution of respondents

    The literal chapter prefix comes from LLM metadata; the item number is
    a real SEQ field so Word keeps it correct and LOT/LOF can collect it.
    """
    paragraph.clear()

    prefix = f"{label} "
    if chapter is not None:
        prefix += f"{chapter}."

    paragraph.add_run(prefix)
    _append_seq_field(paragraph, label, str(index_in_chapter))
    paragraph.add_run(f": {caption_text}")

    # Style + alignment from profile
    caption_cfg = profile.get("captions", {}) or {}
    if role == "table":
        style_name = caption_cfg.get("tableStyleName", "Caption")
        bold = caption_cfg.get("tableBold", True)
    else:
        style_name = caption_cfg.get("figureStyleName", "Caption")
        bold = caption_cfg.get("figureBold", False)

    try:
        paragraph.style = style_name
    except KeyError:
        pass

    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        run.bold = bold


def _build_para_map(doc: Document, blocks: List[Dict[str, Any]]) -> dict:
    """Map P# block IDs -> Paragraph objects using the same index walk as _get_paragraph_by_block_id."""
    result = {}
    para_index = -1
    for block in blocks:
        if block.get("type") != "paragraph":
            continue
        para_index += 1
        if 0 <= para_index < len(doc.paragraphs):
            result[block["id"]] = doc.paragraphs[para_index]
    return result


def _build_table_el_map(doc: Document, blocks: List[Dict[str, Any]]) -> dict:
    """Map T# block IDs -> <w:tbl> XML elements in body order."""
    tbl_elements = [el for el in doc.element.body if el.tag == qn("w:tbl")]
    result = {}
    t_idx = 0
    for block in blocks:
        if block.get("type") == "table":
            if t_idx < len(tbl_elements):
                result[block["id"]] = tbl_elements[t_idx]
            t_idx += 1
    return result


def _reposition(body, caption_el, anchor_el, position: str) -> None:
    """
    Move caption_el to be immediately before/after anchor_el — but only if
    it is currently on the wrong side. Uses lxml addprevious/addnext which
    automatically detaches the element from its current position.
    """
    body_list = list(body)
    try:
        cap_idx = body_list.index(caption_el)
        anc_idx = body_list.index(anchor_el)
    except ValueError:
        return  # element not a direct body child (e.g. inside a table cell)

    if position == "before" and cap_idx > anc_idx:
        anchor_el.addprevious(caption_el)
    elif position == "after" and cap_idx < anc_idx:
        anchor_el.addnext(caption_el)


def _insert_figure_caption(
    doc: Document,
    image_para_el,
    label: str,
    chapter: int,
    index_in_chapter: int,
    caption_text: str,
    profile: Dict[str, Any],
) -> None:
    """
    Create a new caption paragraph immediately after image_para_el.
    Uses doc.add_paragraph() so the paragraph has a proper parent and
    style access, then moves the underlying XML element into position.
    """
    new_para = doc.add_paragraph()          # appended at end; proper parent set
    write_caption_with_seq(
        new_para,
        label=label,
        chapter=chapter,
        index_in_chapter=index_in_chapter,
        caption_text=caption_text,
        profile=profile,
        role="figure",
    )
    # Move XML element to sit right after the image paragraph
    image_para_el.addnext(new_para._p)
    print(f"[INFO] Created figure caption: {label} {chapter}.{index_in_chapter}")


def apply_tables_and_figures(
    doc: Document,
    blocks: List[Dict[str, Any]],
    metadata: Dict[str, Any],
    profile: Dict[str, Any],
) -> None:
    """
    Normalize table and figure numbering & captions using LLM metadata.

    Tables  → caption is placed ABOVE the table.
    Figures → caption is placed BELOW the image paragraph.
              If no caption exists in the original document, one is created.

    MUST run before any prelim pages are inserted (block↔paragraph indices
    are only valid while the document structure is unchanged).
    """
    if "blocks" not in metadata:
        return

    _ensure_caption_style(doc)

    block_meta: Dict[str, Any] = metadata["blocks"]
    table_counters: Dict[int, int] = {}
    figure_counters: Dict[int, int] = {}

    # Pre-build lookup maps once (avoids O(n²) walks per caption)
    para_map = _build_para_map(doc, blocks)
    table_el_map = _build_table_el_map(doc, blocks)

    # Collect position fixes to apply after all SEQ rewrites are done,
    # so the para_map indices stay valid throughout the rewrite phase.
    position_fixes: list = []          # (caption_el, anchor_el, "before"/"after")
    new_figure_captions: list = []     # figures with no existing caption paragraph

    for block in blocks:
        block_id = block.get("id")
        info = block_meta.get(block_id)
        if not info:
            continue

        role = info.get("role")
        if role not in ("table", "figure"):
            continue

        chapter = info.get("chapter") or 1
        info["chapter"] = chapter

        counters = table_counters if role == "table" else figure_counters
        counters[chapter] = counters.get(chapter, 0) + 1
        index_in_chapter = counters[chapter]

        label = "Table" if role == "table" else "Figure"
        caption_text = (info.get("caption") or info.get("rawCaption") or info.get("title") or "Untitled")
        stripped = _strip_existing_number_prefix(caption_text, label)

        caption_block_id = info.get("caption_block_id")

        if not caption_block_id:
            if role == "figure":
                # No caption in the original doc — create one below the image
                parent_id = block.get("parent")
                parent_para = para_map.get(parent_id)
                if parent_para is not None:
                    new_figure_captions.append(
                        (parent_para._p, label, chapter, index_in_chapter, stripped, profile)
                    )
            # Tables without a caption block are left as-is
            continue

        caption_para = para_map.get(caption_block_id)
        if caption_para is None:
            print(f"[WARN] Caption paragraph {caption_block_id} for {block_id} not found")
            continue

        write_caption_with_seq(
            caption_para,
            label=label,
            chapter=chapter,
            index_in_chapter=index_in_chapter,
            caption_text=stripped,
            profile=profile,
            role=role,
        )

        number_key = "tableNumber" if role == "table" else "figureNumber"
        info[number_key] = f"{label} {chapter}.{index_in_chapter}"

        # Schedule the position fix
        if role == "table":
            table_el = table_el_map.get(block_id)
            if table_el is not None:
                # Caption must end up ABOVE the table
                position_fixes.append((caption_para._p, table_el, "before"))
        elif role == "figure":
            parent_id = block.get("parent")
            parent_para = para_map.get(parent_id)
            if parent_para is not None:
                # Caption must end up BELOW the image paragraph
                position_fixes.append((caption_para._p, parent_para._p, "after"))

    # Phase 2: move captions to correct positions
    body = doc.element.body
    for caption_el, anchor_el, pos in position_fixes:
        _reposition(body, caption_el, anchor_el, pos)

    # Phase 3: insert brand-new captions for uncaptioned figures
    for args in new_figure_captions:
        _insert_figure_caption(doc, *args)

    print(
        f"[INFO] Rewrote captions with SEQ fields "
        f"(tables: {sum(table_counters.values())}, figures: {sum(figure_counters.values())})"
    )


def _strip_existing_number_prefix(caption_text: str, label: str) -> str:
    """
    Remove leading 'Table 3.1:' / 'Figure 2 -' style prefixes so we don't
    end up with 'Table 1.1: Table 3.1: ...' after normalization.
    """
    import re

    pattern = rf"^\s*{label}\s*[\dIVXivx]+([.\-]\d+)*\s*[:.\-–]?\s*"
    return re.sub(pattern, "", caption_text, flags=re.IGNORECASE).strip() or caption_text.strip()
