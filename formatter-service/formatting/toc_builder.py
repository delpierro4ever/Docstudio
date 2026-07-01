from typing import Dict, Any, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


def insert_table_of_contents(
    doc: Document,
    metadata: Dict[str, Any],
    anchor_para: Optional[Paragraph] = None,
) -> Optional[Paragraph]:
    """
    Inserts a TABLE OF CONTENTS title and a functional TOC field BEFORE the
    anchor_para, preceded by a simple page break.

    The prelim/main section break is NOT handled here — section_builder
    inserts a real <w:sectPr> paragraph once, at the prelim/main boundary.
    (The old code called add_break(WD_SECTION.NEXT_PAGE), which is invalid:
    add_break only accepts WD_BREAK values and raised on every run.)

    Returns the topmost inserted paragraph (the new anchor for the next
    prelim page inserted above this one).
    """
    toc_meta: Dict[str, Any] = metadata.get("toc", {}) or {}

    if not toc_meta.get("autoInsert", True):
        return None

    # --- 1. PAGE BREAK (separates this page from whatever precedes it) ---
    is_first_paragraph_in_doc = (
        anchor_para is not None
        and doc.paragraphs
        and anchor_para._p is doc.paragraphs[0]._p
    )

    break_para = None
    if anchor_para is not None and not is_first_paragraph_in_doc:
        break_para = anchor_para.insert_paragraph_before()
        break_para.add_run().add_break(WD_BREAK.PAGE)

    # --- 2. INSERT TOC FIELD (just before the anchor) ---
    if anchor_para is not None:
        toc_para = anchor_para.insert_paragraph_before()
    else:
        toc_para = doc.add_paragraph()

    _add_toc_field(toc_para, toc_meta)

    # --- 3. INSERT TOC TITLE (just before the field) ---
    title_text = toc_meta.get("titleText", "TABLE OF CONTENTS")
    title_style_name = toc_meta.get("titleStyleName", "Heading 1")

    title_para = toc_para.insert_paragraph_before(title_text)

    try:
        title_para.style = title_style_name
    except Exception:
        pass

    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Return the TOPMOST paragraph we inserted (the page break when there
    # is one): the next prelim page must be inserted above the break that
    # belongs to THIS page, otherwise the breaks pile up at the front.
    return break_para if break_para is not None else title_para


def _add_toc_field(paragraph, toc_meta: Dict[str, Any]) -> None:
    """
    Inserts a functional TOC field. The result is rendered when the
    field-update step runs (Word's updateFields-on-open flag and/or the
    LibreOffice bake pass — see field_updater.py).
    """
    levels = toc_meta.get("includeHeadingLevels", [1, 2, 3])
    if not levels:
        levels = [1, 2, 3]

    min_level = min(levels)
    max_level = max(levels)
    o_switch = f'\\o "{min_level}-{max_level}"'

    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 1) Begin field: <w:fldChar w:fldCharType="begin" w:dirty="true"/>
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    fld_char_begin.set(qn("w:dirty"), "true")
    run._r.append(fld_char_begin)

    # 2) Instruction text: <w:instrText xml:space="preserve"> TOC ... </w:instrText>
    run = paragraph.add_run()
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = f' TOC {o_switch} \\h \\z \\u '
    run._r.append(instr_text)

    # 3) Separate
    run = paragraph.add_run()
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_char_sep)

    # 4) Placeholder result run (replaced when fields are updated)
    placeholder_run = paragraph.add_run()
    placeholder_run.text = "Right-click and choose 'Update Field' if this text is still visible."
    placeholder_run.italic = True

    # 5) End
    run = paragraph.add_run()
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_end)
