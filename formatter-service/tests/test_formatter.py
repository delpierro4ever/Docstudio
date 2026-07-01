# formatter-service/tests/test_formatter.py

"""
Automated tests for the DOCX formatting pipeline.

These run WITHOUT the LLM (metadata is constructed by hand) and without
LibreOffice baking (DOCSTUDIO_BAKE_FIELDS=0), so they are fast and
deterministic. Run from the formatter-service directory:

    python3 -m unittest discover -s tests -v
"""

import os
import sys
import tempfile
import unittest
from io import BytesIO

# Make formatter-service root importable (config.*, formatting.*, llm.*, docx_parser)
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# Keep tests hermetic: no LLM proofread call, no LibreOffice round-trip
os.environ["DOCSTUDIO_PROOFREAD"] = "0"
os.environ["DOCSTUDIO_BAKE_FIELDS"] = "0"

from docx import Document  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402

from config.profiles import load_profile  # noqa: E402
from docx_parser import DOCXBlockParser  # noqa: E402
from formatting.formatter import format_docx  # noqa: E402
from formatting.abbreviations_builder import extract_abbreviations  # noqa: E402
from llm.proofreader import Proofreader, collect_proofread_candidates  # noqa: E402


def _build_sample_docx(path: str) -> None:
    """A miniature student report with all the structures we format."""
    doc = Document()

    doc.add_paragraph("IMPACT OF ICT ADOPTION ON SMALL BUSINESSES")   # P1 title
    doc.add_paragraph("ABSTRACT")                                     # P2
    doc.add_paragraph(                                                # P3
        "This study examines Information and Communication Technology "
        "(ICT) adoption and its effect on Gross Domestic Product (GDP). "
        "The World Health Organization (WHO) provided some datasets."
    )
    doc.add_paragraph("CHAPTER ONE: INTRODUCTION")                    # P4 (main starts)
    doc.add_paragraph(                                                # P5
        "Small businesses increasingly rely on ICT to compete. "
        "This chapter introduces the study."
    )
    doc.add_paragraph("Table 1: Age distribution of respondents")     # P6 caption
    table = doc.add_table(rows=2, cols=2)                             # T1
    table.cell(0, 0).text = "Age"
    table.cell(0, 1).text = "Count"
    doc.add_paragraph("Figure 3.9: Conceptual framework")             # P7 caption
    doc.add_paragraph("REFERENCES")                                   # P8
    doc.add_paragraph(
        "Doe, J. (2020). ICT and growth. Journal of Development, 4(2), 1-20."
    )                                                                 # P9

    doc.save(path)


def _make_metadata() -> dict:
    """Hand-written classification, as the LLM would produce it."""
    return {
        "blocks": {
            "P1": {"role": "title_page", "section": "prelim"},
            "P2": {"role": "abstract_heading", "section": "prelim"},
            "P3": {"role": "abstract_paragraph", "section": "prelim"},
            "P4": {
                "role": "chapter_heading",
                "section": "main",
                "chapter": 1,
                "headingLevel": 1,
                "title": "CHAPTER ONE: INTRODUCTION",
            },
            "P5": {"role": "body_paragraph", "section": "main"},
            "P6": {"role": "caption", "section": "main"},
            "T1": {
                "role": "table",
                "section": "main",
                "chapter": 1,
                "caption": "Table 1: Age distribution of respondents",
                "caption_block_id": "P6",
            },
            "P7": {"role": "caption", "section": "main"},
            "F1": {
                "role": "figure",
                "section": "main",
                "chapter": 1,
                "caption": "Figure 3.9: Conceptual framework",
                "caption_block_id": "P7",
            },
            "P8": {"role": "references_heading", "section": "main"},
            "P9": {"role": "reference_entry", "section": "main"},
        },
        "sections": {"prelim_ends_before_block_id": "P4"},
        "toc": {"autoInsert": True, "includeHeadingLevels": [1, 2, 3]},
    }


class FormatterPipelineTest(unittest.TestCase):
    """End-to-end format_docx run on a synthetic document."""

    @classmethod
    def setUpClass(cls):
        cls.tmp = tempfile.mkdtemp(prefix="docstudio_test_")
        cls.input_path = os.path.join(cls.tmp, "sample.docx")
        _build_sample_docx(cls.input_path)

        parser = DOCXBlockParser(cls.input_path)
        cls.blocks = parser.parse()

        # The synthetic figure is text-only (no image), so add a fake
        # figure block referencing the caption, mirroring real parses.
        if not any(b["id"].startswith("F") for b in cls.blocks):
            cls.blocks.append({"id": "F1", "type": "image", "rel": None})

        cls.metadata = _make_metadata()
        cls.output_bytes = format_docx(
            input_path=cls.input_path,
            blocks=cls.blocks,
            metadata=cls.metadata,
            profile_id="ub-v1",
        )
        cls.out_doc = Document(BytesIO(cls.output_bytes))
        cls.doc_xml = cls.out_doc.element.xml

    # -- basic sanity ------------------------------------------------------

    def test_pipeline_produces_docx(self):
        self.assertGreater(len(self.output_bytes), 0)
        self.assertTrue(self.output_bytes.startswith(b"PK"))  # zip magic

    def test_no_placeholder_crash_left_lof_missing(self):
        """The old WD_SECTION.NEXT_PAGE crash killed LIST OF FIGURES."""
        texts = [p.text.strip() for p in self.out_doc.paragraphs]
        self.assertIn("LIST OF FIGURES", texts)
        self.assertIn("LIST OF TABLES", texts)
        self.assertIn("TABLE OF CONTENTS", texts)

    # -- field-update step -------------------------------------------------

    def test_update_fields_flag_set(self):
        settings_xml = self.out_doc.settings.element.xml
        self.assertIn("updateFields", settings_xml)

    def test_toc_field_present(self):
        self.assertIn("TOC \\o", self.doc_xml.replace("&quot;", '"'))

    def test_lot_lof_fields_reference_seq_labels(self):
        xml = self.doc_xml.replace("&quot;", '"')
        self.assertIn('TOC \\h \\z \\c "Table"', xml)
        self.assertIn('TOC \\h \\z \\c "Figure"', xml)

    # -- SEQ captions --------------------------------------------------------

    def test_captions_rewritten_with_seq_fields(self):
        xml = self.doc_xml
        self.assertIn("SEQ Table", xml)
        self.assertIn("SEQ Figure", xml)

    def test_caption_number_normalized(self):
        texts = [p.text for p in self.out_doc.paragraphs]
        table_caps = [t for t in texts if t.startswith("Table 1.")]
        self.assertTrue(table_caps, f"no normalized table caption in {texts}")
        # Old prefix ("Table 1:" / "Figure 3.9:") must not survive inside
        self.assertIn(": Age distribution of respondents", table_caps[0])
        fig_caps = [t for t in texts if t.startswith("Figure 1.")]
        self.assertTrue(fig_caps, f"no normalized figure caption in {texts}")
        self.assertNotIn("3.9", fig_caps[0])

    # -- section break + page numbering -------------------------------------

    def test_two_sections_created(self):
        self.assertGreaterEqual(len(self.out_doc.sections), 2)

    def test_page_number_formats(self):
        sections = self.out_doc.sections
        prelim_pgnum = sections[0]._sectPr.find(qn("w:pgNumType"))
        main_pgnum = sections[1]._sectPr.find(qn("w:pgNumType"))
        self.assertIsNotNone(prelim_pgnum)
        self.assertIsNotNone(main_pgnum)
        self.assertEqual(prelim_pgnum.get(qn("w:fmt")), "lowerRoman")
        self.assertEqual(main_pgnum.get(qn("w:fmt")), "decimal")
        self.assertEqual(main_pgnum.get(qn("w:start")), "1")

    def test_section_break_lands_at_chapter_boundary(self):
        """The paragraph carrying the prelim sectPr must sit immediately
        before CHAPTER ONE — the old code put it pages too early."""
        # heading normalization rewrites "CHAPTER ONE: ..." to "CHAPTER 1: ..."
        paras = self.out_doc.paragraphs
        chapter_idx = next(
            i for i, p in enumerate(paras) if p.text.strip().startswith("CHAPTER")
        )
        sect_para = paras[chapter_idx - 1]
        pPr = sect_para._p.find(qn("w:pPr"))
        self.assertIsNotNone(pPr, "no pPr on paragraph before CHAPTER ONE")
        self.assertIsNotNone(
            pPr.find(qn("w:sectPr")),
            "section break is not immediately before the chapter heading",
        )

    def test_prelim_pages_in_required_order(self):
        texts = [p.text.strip() for p in self.out_doc.paragraphs]
        order = [
            texts.index("TABLE OF CONTENTS"),
            texts.index("LIST OF TABLES"),
            texts.index("LIST OF FIGURES"),
            texts.index("LIST OF ABBREVIATIONS"),
            next(i for i, t in enumerate(texts) if t.startswith("CHAPTER")),
        ]
        self.assertEqual(order, sorted(order))

    def test_page_breaks_between_generated_pages(self):
        """Each generated prelim page after the first must be preceded by a
        page break (the old anchor logic piled all breaks at the front)."""
        paras = self.out_doc.paragraphs
        texts = [p.text.strip() for p in paras]
        for title in ("LIST OF TABLES", "LIST OF FIGURES", "LIST OF ABBREVIATIONS"):
            idx = texts.index(title)
            prev_xml = paras[idx - 1]._p.xml
            self.assertIn(
                'w:type="page"', prev_xml,
                f"no page break immediately before {title}",
            )

    # -- abbreviations -------------------------------------------------------

    def test_abbreviations_page_generated(self):
        texts = [p.text.strip() for p in self.out_doc.paragraphs]
        self.assertIn("LIST OF ABBREVIATIONS", texts)
        joined = "\n".join(texts)
        self.assertIn("ICT", joined)
        self.assertIn("Information and Communication Technology", joined)

    # -- footers -------------------------------------------------------------

    def test_footers_have_page_fields(self):
        for i, section in enumerate(self.out_doc.sections[:2]):
            footer_xml = section.footer._element.xml
            self.assertIn("PAGE", footer_xml, f"section {i} footer lacks PAGE field")


class AbbreviationExtractionTest(unittest.TestCase):
    def test_definition_and_bare_acronyms(self):
        blocks = [
            {"id": "P1", "type": "paragraph",
             "text": "The World Bank (WB) and the Gross Domestic Product (GDP) matter. HIV rates fell."},
            {"id": "P2", "type": "paragraph", "text": "CHAPTER ONE"},
        ]
        metadata = {"blocks": {"P2": {"role": "chapter_heading"}}}
        result = dict(extract_abbreviations(blocks, metadata))
        self.assertEqual(result.get("GDP"), "Gross Domestic Product")
        self.assertIn("HIV", result)
        self.assertNotIn("CHAPTER", result)

    def test_stopwords_excluded(self):
        blocks = [{"id": "P1", "type": "paragraph",
                   "text": "LIST OF TABLES AND THE ABSTRACT"}]
        result = dict(extract_abbreviations(blocks, {}))
        self.assertEqual(result, {})


class ProofreaderSanitizeTest(unittest.TestCase):
    def test_keeps_only_valid_corrections(self):
        batch = {"P1": "Ths is a sentnce.", "P2": "All good here."}
        raw = {
            "P1": "This is a sentence.",
            "P2": "All good here.",          # unchanged → dropped
            "P9": "Invented paragraph.",     # unknown id → dropped
            "P3": 42,                        # wrong type → dropped
        }
        clean = Proofreader._sanitize(raw, batch)
        self.assertEqual(clean, {"P1": "This is a sentence."})

    def test_rejects_runaway_rewrites(self):
        batch = {"P1": "Short text here."}
        raw = {"P1": "An enormously long rewrite " * 20}
        clean = Proofreader._sanitize(raw, batch)
        self.assertEqual(clean, {})

    def test_candidate_selection_skips_headings_and_short_text(self):
        blocks = [
            {"id": "P1", "type": "paragraph", "text": "CHAPTER ONE: INTRODUCTION"},
            {"id": "P2", "type": "paragraph", "text": "Short."},
            {"id": "P3", "type": "paragraph",
             "text": "This paragraph is long enough to be worth proofreading."},
        ]
        metadata = {"blocks": {
            "P1": {"role": "chapter_heading"},
            "P2": {"role": "body_paragraph"},
            "P3": {"role": "body_paragraph"},
        }}
        candidates = collect_proofread_candidates(blocks, metadata)
        self.assertEqual(list(candidates.keys()), ["P3"])


class ProfileTest(unittest.TestCase):
    def test_default_profile_has_13pt_subheadings(self):
        profile = load_profile("ub-v1")
        self.assertEqual(profile["headings"]["subHeading1"]["fontSize"], 13)
        self.assertEqual(profile["headings"]["subHeading2"]["fontSize"], 13)

    def test_large_print_profile_exists(self):
        profile = load_profile("ub-v2")
        self.assertEqual(profile["font"]["size"], 14)
        self.assertEqual(profile["paragraph"]["lineSpacing"], 2.0)

    def test_unknown_profile_falls_back_to_default(self):
        profile = load_profile("does-not-exist")
        self.assertEqual(profile["id"], "ub-v1")


if __name__ == "__main__":
    unittest.main(verbosity=2)
