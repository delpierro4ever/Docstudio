
import docx
from docx.shared import Pt
import os

def get_style_info(paragraph):
    style = paragraph.style
    font_name = None
    font_size = None
    if style and hasattr(style, 'font'):
        font_name = style.font.name
        font_size = style.font.size
    
    # Check run overrides
    for run in paragraph.runs:
        if run.font.name:
            font_name = run.font.name
        if run.font.size:
            font_size = run.font.size
    
    return {
        "alignment": paragraph.alignment,
        "font_name": font_name,
        "font_size": font_size
    }

def analyze_doc(path):
    print(f"Analyzing {path}...")
    try:
        doc = docx.Document(path)
    except Exception as e:
        print(f"Error reading {path}: {e}")
        return

    print(f"Total Paragraphs: {len(doc.paragraphs)}")
    
    # Analyze sections (margins)
    for i, section in enumerate(doc.sections):
        print(f"Section {i} Margins: L={section.left_margin}, R={section.right_margin}, T={section.top_margin}, B={section.bottom_margin}")
        print(f"Section {i} Page Size: W={section.page_width}, H={section.page_height}")
    
    # Analyze first 10 paragraphs styles
    print("First 10 Paragraph Styles:")
    for i, p in enumerate(doc.paragraphs[:10]):
        info = get_style_info(p)
        print(f"P{i}: Text='{p.text[:30]}...' Style={p.style.name} Font={info['font_name']} Size={info['font_size']} Align={info['alignment']}")

    # Analyze Headers/Footers
    print("Headers/Footers:")
    for i, section in enumerate(doc.sections):
        header = section.header
        footer = section.footer
        if header:
            print(f"Section {i} Header: {[p.text for p in header.paragraphs]}")
        if footer:
            print(f"Section {i} Footer: {[p.text for p in footer.paragraphs]}")


    # Analyze Line Spacing
    if len(doc.paragraphs) > 0:
        p = doc.paragraphs[0]
        print(f"Line Spacing: {p.paragraph_format.line_spacing}, Rule: {p.paragraph_format.line_spacing_rule}")


print("\n=== SAMPLE EDITED (EXPECTED) ===")
analyze_doc("examples/sample_edited.docx")
# print("\n=== LATEST RESULT ===")
# analyze_doc("examples/latest_result.docx")


