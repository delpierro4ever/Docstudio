"""
Analyze differences between result.docx (actual) and sample_edited.docx (expected)
"""
from docx import Document
import os

def analyze_document(filepath):
    """Extract key information from a document"""
    doc = Document(filepath)
    
    info = {
        'filename': os.path.basename(filepath),
        'total_paragraphs': len(doc.paragraphs),
        'total_sections': len(doc.sections),
        'paragraphs': []
    }
    
    # Analyze first 50 paragraphs in detail
    for i, para in enumerate(doc.paragraphs[:50]):
        para_info = {
            'index': i,
            'text': para.text[:100] if para.text else '',  # First 100 chars
            'style': para.style.name if para.style else 'None',
            'alignment': str(para.alignment) if para.alignment else 'None',
        }
        
        # Check for page breaks
        if para._element.xpath('.//w:br[@w:type="page"]'):
            para_info['has_page_break'] = True
            
        info['paragraphs'].append(para_info)
    
    # Section information
    info['sections_info'] = []
    for i, section in enumerate(doc.sections):
        section_info = {
            'index': i,
            'start_type': str(section.start_type),
        }
        
        # Try to get page numbering info
        try:
            sectPr = section._sectPr
            from docx.oxml.ns import qn
            pgNumType = sectPr.find(qn('w:pgNumType'))
            if pgNumType is not None:
                section_info['page_num_format'] = pgNumType.get(qn('w:fmt'))
                section_info['page_num_start'] = pgNumType.get(qn('w:start'))
        except:
            pass
            
        info['sections_info'].append(section_info)
    
    return info

def compare_documents():
    base_path = r"C:\Users\Cephas\Documents\Programming\Student-Report-Formatter\examples"
    
    print("="*80)
    print("DOCUMENT COMPARISON ANALYSIS")
    print("="*80)
    
    # Analyze result.docx (current output)
    print("\n### RESULT.DOCX (Current Pipeline Output) ###\n")
    result_info = analyze_document(os.path.join(base_path, "result.docx"))
    print(f"Total Paragraphs: {result_info['total_paragraphs']}")
    print(f"Total Sections: {result_info['total_sections']}")
    print("\nFirst 30 Paragraphs:")
    for p in result_info['paragraphs'][:30]:
        print(f"  [{p['index']:3d}] Style: {p['style']:30s} | {p['text']}")
        if p.get('has_page_break'):
            print(f"       >>> PAGE BREAK HERE <<<")
    
    print("\nSection Information:")
    for s in result_info['sections_info']:
        print(f"  Section {s['index']}: start_type={s['start_type']}, "
              f"page_fmt={s.get('page_num_format', 'N/A')}, "
              f"page_start={s.get('page_num_start', 'N/A')}")
    
    # Analyze sample_edited.docx (expected output)
    print("\n" + "="*80)
    print("### SAMPLE_EDITED.DOCX (Expected Output) ###\n")
    expected_info = analyze_document(os.path.join(base_path, "sample_edited.docx"))
    print(f"Total Paragraphs: {expected_info['total_paragraphs']}")
    print(f"Total Sections: {expected_info['total_sections']}")
    print("\nFirst 30 Paragraphs:")
    for p in expected_info['paragraphs'][:30]:
        print(f"  [{p['index']:3d}] Style: {p['style']:30s} | {p['text']}")
        if p.get('has_page_break'):
            print(f"       >>> PAGE BREAK HERE <<<")
    
    print("\nSection Information:")
    for s in expected_info['sections_info']:
        print(f"  Section {s['index']}: start_type={s['start_type']}, "
              f"page_fmt={s.get('page_num_format', 'N/A')}, "
              f"page_start={s.get('page_num_start', 'N/A')}")
    
    # Key Differences
    print("\n" + "="*80)
    print("### KEY DIFFERENCES ###\n")
    print(f"Paragraph count difference: {expected_info['total_paragraphs'] - result_info['total_paragraphs']}")
    print(f"Section count difference: {expected_info['total_sections'] - result_info['total_sections']}")
    
    # Compare styles
    print("\nStyle differences in first 30 paragraphs:")
    for i in range(min(30, len(result_info['paragraphs']), len(expected_info['paragraphs']))):
        result_style = result_info['paragraphs'][i]['style']
        expected_style = expected_info['paragraphs'][i]['style']
        if result_style != expected_style:
            print(f"  Para {i}: Result='{result_style}' vs Expected='{expected_style}'")
            print(f"    Text: {expected_info['paragraphs'][i]['text'][:80]}")

if __name__ == "__main__":
    compare_documents()
