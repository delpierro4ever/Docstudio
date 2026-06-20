"""
Compare sample.docx (input), sample_edited.docx (expected output), 
and latest_result.docx (actual output) to understand differences.
"""
from docx import Document
from docx.oxml.ns import qn
import os
import json

EXAMPLES_DIR = os.path.join(os.path.dirname(__file__), "..", "examples")

def analyze_doc(filepath):
    """Get detailed structural info about a docx file."""
    doc = Document(filepath)
    info = {
        "filename": os.path.basename(filepath),
        "total_paragraphs": len(doc.paragraphs),
        "total_sections": len(doc.sections),
        "sections": [],
        "paragraphs": [],
    }
    
    # Sections info
    for si, section in enumerate(doc.sections):
        sect_info = {
            "index": si,
            "start_type": str(section.start_type) if section.start_type else "None",
            "page_width": str(section.page_width),
            "page_height": str(section.page_height),
        }
        
        # Check page numbering
        sectPr = section._sectPr
        pgNumType = sectPr.find(qn('w:pgNumType'))
        if pgNumType is not None:
            sect_info["pgNumType_fmt"] = pgNumType.get(qn('w:fmt'))
            sect_info["pgNumType_start"] = pgNumType.get(qn('w:start'))
        
        # Check footer
        try:
            footer = section.footer
            sect_info["has_footer"] = not footer.is_linked_to_previous if footer else False
            if footer and footer.paragraphs:
                sect_info["footer_text"] = [p.text for p in footer.paragraphs]
        except:
            sect_info["has_footer"] = False
            
        # Check header
        try:
            header = section.header
            sect_info["has_header"] = not header.is_linked_to_previous if header else False
            if header and header.paragraphs:
                sect_info["header_text"] = [p.text for p in header.paragraphs]
        except:
            sect_info["has_header"] = False
            
        info["sections"].append(sect_info)
    
    # Paragraph info (first 150 and last 30)
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if i < 150 or i > len(doc.paragraphs) - 30 or text:
            para_info = {
                "idx": i,
                "text": text[:120] if text else "(empty)",
                "style": p.style.name,
                "alignment": str(p.alignment) if p.alignment else "None",
            }
            
            # Check for bold
            if p.runs:
                bold_runs = sum(1 for r in p.runs if r.bold)
                para_info["bold_ratio"] = f"{bold_runs}/{len(p.runs)}"
                
                # Font info from first run
                first_run = p.runs[0]
                para_info["font_name"] = first_run.font.name
                para_info["font_size"] = str(first_run.font.size) if first_run.font.size else "None"
            
            # Check for section break in paragraph properties
            pPr = p._p.find(qn('w:pPr'))
            if pPr is not None:
                sectPr = pPr.find(qn('w:sectPr'))
                if sectPr is not None:
                    para_info["has_section_break"] = True
                    type_elem = sectPr.find(qn('w:type'))
                    if type_elem is not None:
                        para_info["section_break_type"] = type_elem.get(qn('w:val'))
            
            # Check for page breaks
            for r in p.runs:
                if 'w:br' in r._element.xml:
                    if 'type="page"' in r._element.xml:
                        para_info["has_page_break"] = True
                        
            # Check for TOC field
            for r in p.runs:
                fld_chars = r._element.findall(qn('w:fldChar'))
                instr_texts = r._element.findall(qn('w:instrText'))
                if fld_chars or instr_texts:
                    para_info["has_field"] = True
                    for it in instr_texts:
                        para_info["field_text"] = it.text
            
            info["paragraphs"].append(para_info)
    
    return info


def main():
    files = [
        os.path.join(EXAMPLES_DIR, "sample.docx"),
        os.path.join(EXAMPLES_DIR, "sample_edited.docx"),
        os.path.join(EXAMPLES_DIR, "latest_result.docx"),
        os.path.join(EXAMPLES_DIR, "result.docx"),
    ]
    
    for f in files:
        if os.path.exists(f):
            print(f"\n{'='*80}")
            print(f"ANALYZING: {os.path.basename(f)}")
            print(f"{'='*80}")
            info = analyze_doc(f)
            
            print(f"\nTotal paragraphs: {info['total_paragraphs']}")
            print(f"Total sections: {info['total_sections']}")
            
            print(f"\n--- SECTIONS ---")
            for s in info['sections']:
                print(f"  Section {s['index']}: {json.dumps(s, indent=4, default=str)}")
            
            print(f"\n--- KEY PARAGRAPHS (non-empty or first 150) ---")
            for p in info['paragraphs']:
                text = p['text']
                if text == "(empty)":
                    continue
                line = f"  [{p['idx']:3d}] [{p['style']:20s}] "
                if p.get('has_section_break'):
                    line += "[SECT-BRK] "
                if p.get('has_page_break'):
                    line += "[PG-BRK] "
                if p.get('has_field'):
                    line += f"[FIELD:{p.get('field_text', '?')[:30]}] "
                line += f"{text[:100]}"
                print(line)
        else:
            print(f"File not found: {f}")


if __name__ == "__main__":
    main()
