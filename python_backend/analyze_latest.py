"""
Comprehensive analysis of the latest output
"""
from docx import Document
from docx.oxml.ns import qn
import os

output_file = "downloaded_8c4b6b46-0ad7-41ee-bc64-95d3082c4a28.docx"

if os.path.exists(output_file):
    doc = Document(output_file)
    
    print(f"Total Paragraphs: {len(doc.paragraphs)}")
    print(f"Total Sections: {len(doc.sections)}")
    
    print("\n### Heading 1 Paragraphs ###")
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == "Heading 1":
            text = p.text[:80] if p.text else ""
            print(f"  [{i:3d}] {text}")
    
    print("\n### First 50 Paragraphs with Styles ###")
    for i, p in enumerate(doc.paragraphs[:50]):
        text = p.text[:80] if p.text else ""
        if text:  # Only show non-empty
            print(f"  [{i:3d}] {p.style.name:20s} | {text}")
    
    print("\n### Section Info ###")
    for i, section in enumerate(doc.sections):
        sectPr = section._sectPr
        pgNumType = sectPr.find(qn('w:pgNumType'))
        if pgNumType is not None:
            fmt = pgNumType.get(qn('w:fmt'))
            start = pgNumType.get(qn('w:start'))
            print(f"Section {i}: fmt={fmt}, start={start}")
        else:
            print(f"Section {i}: No page numbering set")
else:
    print(f"File not found: {output_file}")
