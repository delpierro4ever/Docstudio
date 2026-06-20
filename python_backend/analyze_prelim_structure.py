"""
Analyze sample_edited.docx to understand the structure of preliminary sections
"""
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

edited_path = r"C:\Users\Cephas\Documents\Programming\Student-Report-Formatter\examples\sample_edited.docx"

if os.path.exists(edited_path):
    doc = Document(edited_path)
    
    print("SAMPLE_EDITED.DOCX STRUCTURE ANALYSIS")
    print("="*80)
    
    # Find each preliminary section and analyze its structure
    sections_to_find = {
        "Certification": None,
        "Dedication": None,
        "Acknowledgement": None,
        "Abstract": None,
        "List of figures": None,
        "List of Tables": None,
        "List of Abbreviations": None,
        "Chapter 1": None,
        "References": None
    }
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        for section_name in sections_to_find:
            if section_name.upper() in text.upper() and sections_to_find[section_name] is None:
                sections_to_find[section_name] = i
    
    print("\n### SECTION START INDICES ###")
    for name, idx in sections_to_find.items():
        print(f"{name:25s}: {idx if idx is not None else 'NOT FOUND'}")
    
    # Analyze each preliminary section in detail
    print("\n### PRELIMINARY SECTION DETAILS ###")
    
    for section_name, start_idx in sections_to_find.items():
        if start_idx is None or "Chapter" in section_name or "References" in section_name:
            continue
            
        # Find the next section
        next_idx = len(doc.paragraphs)
        for name, idx in sections_to_find.items():
            if idx is not None and idx > start_idx and idx < next_idx:
                next_idx = idx
        
        print(f"\n{section_name} (paragraphs {start_idx} to {next_idx-1}):")
        print(f"  Total paragraphs: {next_idx - start_idx}")
        
        # Show first few paragraphs
        for i in range(start_idx, min(start_idx + 5, next_idx)):
            p = doc.paragraphs[i]
            text = p.text[:100] if p.text else ""
            alignment = "CENTER" if p.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER else "LEFT"
            print(f"    [{i:3d}] {p.style.name:20s} {alignment:8s} | {text}")
    
    # Check abbreviations section specifically
    print("\n### ABBREVIATIONS SECTION DETAIL ###")
    abbrev_idx = sections_to_find.get("List of Abbreviations")
    if abbrev_idx:
        chapter_idx = sections_to_find.get("Chapter 1", len(doc.paragraphs))
        print(f"Abbreviations section: paragraphs {abbrev_idx} to {chapter_idx-1}")
        for i in range(abbrev_idx, min(abbrev_idx + 25, chapter_idx)):
            p = doc.paragraphs[i]
            text = p.text[:100] if p.text else ""
            if text:
                print(f"  [{i:3d}] {p.style.name:20s} | {text}")
    
else:
    print(f"File not found: {edited_path}")
