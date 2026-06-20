"""
Analyze the original sample.docx to see what sections it has
"""
from docx import Document
import os

sample_path = r"C:\Users\Cephas\Documents\Programming\Student-Report-Formatter\examples\sample.docx"

if os.path.exists(sample_path):
    doc = Document(sample_path)
    
    print(f"Total Paragraphs: {len(doc.paragraphs)}")
    
    print("\n### ALL Paragraphs (showing first 100 chars of each) ###")
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        style = p.style.name
        if text:  # Only show non-empty
            print(f"  [{i:3d}] {style:20s} | {text[:100]}")
else:
    print(f"File not found: {sample_path}")
