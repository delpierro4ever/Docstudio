"""
Check what's actually in the original sample.docx input file
"""
from docx import Document
import os

sample_path = r"C:\Users\Cephas\Documents\Programming\Student-Report-Formatter\examples\sample.docx"

if os.path.exists(sample_path):
    doc = Document(sample_path)
    
    print(f"ORIGINAL INPUT: sample.docx")
    print(f"{'='*80}\n")
    
    print(f"Total Paragraphs: {len(doc.paragraphs)}")
    print(f"Total Sections: {len(doc.sections)}")
    
    print("\n### ALL NON-EMPTY PARAGRAPHS ###")
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text:
            print(f"  [{i:3d}] {p.style.name:20s} | {text[:100]}")
    
    print("\n### SEARCH FOR KEY SECTIONS ###")
    keywords = ["CERTIFICATION", "DEDICATION", "ACKNOWLEDGEMENT", "ABSTRACT", "ABBREVIATION", "REFERENCE"]
    for keyword in keywords:
        found = False
        for i, p in enumerate(doc.paragraphs):
            if keyword in p.text.upper():
                print(f"  {keyword}: Found at index {i} - {p.text[:80]}")
                found = True
                break
        if not found:
            print(f"  {keyword}: NOT FOUND")
else:
    print(f"File not found: {sample_path}")
