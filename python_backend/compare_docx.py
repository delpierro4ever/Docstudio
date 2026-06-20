import docx
import os
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def get_paragraph_properties(para, index):
    props = {
        "index": index,
        "text_preview": para.text[:30] + "..." if len(para.text) > 30 else para.text,
        "style": para.style.name,
        "alignment": para.alignment,
        "runs": []
    }
    
    # Check runs for specific formatting overrides
    for i, run in enumerate(para.runs):
        run_props = {
            "font_name": run.font.name,
            "font_size": run.font.size.pt if run.font.size else None,
            "bold": run.bold,
            "italic": run.italic,
            "text": run.text[:20]
        }
        props["runs"].append(run_props)
    return props

def compare_file_properties(result_path, sample_edited_path):
    print(f"Comparing:")
    print(f"  Current Output: {os.path.basename(result_path)}")
    print(f"  Expected Output: {os.path.basename(sample_edited_path)}")
    print("="*60)

    try:
        doc_result = docx.Document(result_path)
        doc_expected = docx.Document(sample_edited_path)
    except Exception as e:
        print(f"Error opening files: {e}")
        return

    # Compare Section Properties (Margins, etc.)
    print("\n--- Section Properties (First Section) ---")
    
    def print_section_props(doc, label):
        section = doc.sections[0]
        try:
            header_distance = section.header_distance.pt if section.header_distance else "Default"
            footer_distance = section.footer_distance.pt if section.footer_distance else "Default"
            print(f"{label}:")
            print(f"  Header Distance: {header_distance}")
            print(f"  Footer Distance: {footer_distance}")
            print(f"  Start Number: {section.start_type}")
        except Exception as e:
            print(f"  Error reading section props: {e}")

    print_section_props(doc_result, "Current Output")
    print_section_props(doc_expected, "Expected Output")

    # Compare Headers/Footers
    print("\n--- Headers & Footers Content (First Section) ---")
    def get_header_footer_content(doc):
        section = doc.sections[0]
        header_text = [p.text for p in section.header.paragraphs if p.text.strip()]
        footer_text = [p.text for p in section.footer.paragraphs if p.text.strip()]
        return header_text, footer_text

    h_res, f_res = get_header_footer_content(doc_result)
    h_exp, f_exp = get_header_footer_content(doc_expected)
    
    print(f"Current Output Header: {h_res}")
    print(f"Expected Output Header: {h_exp}")
    print(f"Current Output Footer: {f_res}")
    print(f"Expected Output Footer: {f_exp}")

    # Compare Paragraphs (checking for global font/style consistency)
    print("\n--- Paragraph Formatting Sample (First 5 Non-Empty) ---")
    
    def analyze_first_paragraphs(doc):
        samples = []
        count = 0
        for p in doc.paragraphs:
            if not p.text.strip():
                continue
            
            font_names = set()
            font_sizes = set()
            for r in p.runs:
                if r.font.name: font_names.add(r.font.name)
                if r.font.size: font_sizes.add(r.font.size.pt)
            
            samples.append({
                "text": p.text[:40],
                "style": p.style.name,
                "fonts": list(font_names) if font_names else ["Inherited/None"],
                "sizes": list(font_sizes) if font_sizes else ["Inherited/None"]
            })
            count += 1
            if count >= 5: break
        return samples

    res_samples = analyze_first_paragraphs(doc_result)
    exp_samples = analyze_first_paragraphs(doc_expected)

    print("\nCurrent Output Samples:")
    for s in res_samples: print(f"  {s}")
    
    print("\nExpected Output Samples:")
    for s in exp_samples: print(f"  {s}")

if __name__ == "__main__":
    base_dir = os.path.dirname(os.path.dirname(__file__))
    result_path = os.path.join(base_dir, "examples", "result.docx")
    sample_edited_path = os.path.join(base_dir, "examples", "sample_edited.docx")
    
    compare_file_properties(result_path, sample_edited_path)
