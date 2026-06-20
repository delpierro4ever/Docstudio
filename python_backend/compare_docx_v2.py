import docx
import os
from docx.shared import Pt
import json

def analyze_document(file_path):
    try:
        doc = docx.Document(file_path)
    except Exception as e:
        return {"error": str(e)}

    analysis = {
        "filename": os.path.basename(file_path),
        "section_properties": {},
        "header_footer": {},
        "paragraph_samples": []
    }

    # Section Props
    try:
        section = doc.sections[0]
        analysis["section_properties"] = {
            "header_distance": section.header_distance.pt if section.header_distance else "Default",
            "footer_distance": section.footer_distance.pt if section.footer_distance else "Default",
            "start_type": str(section.start_type)
        }
        
        # Header/Footer content
        h_text = [p.text for p in section.header.paragraphs if p.text.strip()]
        f_text = [p.text for p in section.footer.paragraphs if p.text.strip()]
        analysis["header_footer"] = {
            "header_content": h_text,
            "footer_content": f_text
        }
        
    except Exception as e:
        analysis["section_error"] = str(e)

    # Paragraph Samples (first 10)
    count = 0
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        
        font_names = set()
        font_sizes = set()
        for r in p.runs:
            if r.font.name: font_names.add(r.font.name)
            if r.font.size: font_sizes.add(r.font.size.pt)
        
        analysis["paragraph_samples"].append({
            "text_start": p.text[:30],
            "style": p.style.name,
            "fonts": list(font_names) if font_names else ["Inherited"],
            "sizes": list(font_sizes) if font_sizes else ["Inherited"],
            "alignment": str(p.alignment)
        })
        count += 1
        if count >= 10: break
    
    return analysis

if __name__ == "__main__":
    base_dir = os.path.dirname(os.path.dirname(__file__))
    result_path = os.path.join(base_dir, "examples", "result.docx")
    sample_edited_path = os.path.join(base_dir, "examples", "sample_edited.docx")
    
    res_analysis = analyze_document(result_path)
    exp_analysis = analyze_document(sample_edited_path)
    
    report = {
        "current_result": res_analysis,
        "expected_result": exp_analysis
    }
    
    # Write to a clean text file
    with open("detailed_comparison.json", "w", encoding="utf-8") as f:
        json.dump(report, f, indent=2)
    
    print("Comparison complete. Saved to detailed_comparison.json")
