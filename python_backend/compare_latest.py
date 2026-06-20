"""
Detailed comparison between latest_result.docx and sample_edited.docx
"""
from docx import Document
from docx.oxml.ns import qn
import os

def analyze_doc(filepath, label):
    """Analyze a document in detail"""
    if not os.path.exists(filepath):
        print(f"ERROR: {filepath} not found")
        return None
    
    doc = Document(filepath)
    
    print(f"\n{'='*80}")
    print(f"{label}: {os.path.basename(filepath)}")
    print(f"{'='*80}\n")
    
    print(f"Total Paragraphs: {len(doc.paragraphs)}")
    print(f"Total Sections: {len(doc.sections)}")
    print(f"Total Tables: {len(doc.tables)}")
    
    # Section details
    print("\n### SECTION DETAILS ###")
    for i, section in enumerate(doc.sections):
        sectPr = section._sectPr
        pgNumType = sectPr.find(qn('w:pgNumType'))
        
        print(f"\nSection {i}:")
        print(f"  Start type: {section.start_type}")
        
        if pgNumType is not None:
            fmt = pgNumType.get(qn('w:fmt'))
            start = pgNumType.get(qn('w:start'))
            print(f"  Page numbering: format={fmt}, start={start}")
        else:
            print(f"  Page numbering: Not set")
    
    # Heading 1 paragraphs
    print("\n### HEADING 1 PARAGRAPHS ###")
    h1_count = 0
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == "Heading 1":
            h1_count += 1
            text = p.text[:80] if p.text else ""
            print(f"  [{i:3d}] {text}")
    print(f"Total Heading 1: {h1_count}")
    
    # First 60 paragraphs with styles
    print("\n### FIRST 60 PARAGRAPHS ###")
    for i, p in enumerate(doc.paragraphs[:60]):
        text = p.text[:80] if p.text else ""
        if text or p.style.name.startswith('Heading'):
            print(f"  [{i:3d}] {p.style.name:20s} | {text}")
    
    return doc

# Compare both documents
base_path = r"C:\Users\Cephas\Documents\Programming\Student-Report-Formatter"

# Find latest_result.docx
latest_path = None
for root, dirs, files in os.walk(base_path):
    for file in files:
        if file.startswith("latest_result") and file.endswith(".docx"):
            latest_path = os.path.join(root, file)
            break
    if latest_path:
        break

if not latest_path:
    print("ERROR: latest_result.docx not found!")
    print("Searching in common locations...")
    # Try common locations
    possible_paths = [
        os.path.join(base_path, "latest_result.docx"),
        os.path.join(base_path, "examples", "latest_result.docx"),
        os.path.join(base_path, "python_backend", "latest_result.docx"),
    ]
    for path in possible_paths:
        if os.path.exists(path):
            latest_path = path
            break

expected_path = os.path.join(base_path, "examples", "sample_edited.docx")

if latest_path:
    latest_doc = analyze_doc(latest_path, "LATEST RESULT")
else:
    print("Could not find latest_result.docx")
    
expected_doc = analyze_doc(expected_path, "EXPECTED RESULT")

# Key differences
if latest_path and latest_doc and expected_doc:
    print(f"\n{'='*80}")
    print("KEY DIFFERENCES")
    print(f"{'='*80}\n")
    
    print(f"Paragraph count: Latest={len(latest_doc.paragraphs)}, Expected={len(expected_doc.paragraphs)}, Diff={len(expected_doc.paragraphs) - len(latest_doc.paragraphs)}")
    print(f"Section count: Latest={len(latest_doc.sections)}, Expected={len(expected_doc.sections)}")
    print(f"Table count: Latest={len(latest_doc.tables)}, Expected={len(expected_doc.tables)}")
