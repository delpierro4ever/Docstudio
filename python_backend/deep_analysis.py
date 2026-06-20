"""
Deep analysis of formatting differences
"""
from docx import Document
import os

def deep_analyze(filepath):
    """Deep analysis of document structure"""
    doc = Document(filepath)
    
    print(f"\n{'='*80}")
    print(f"DEEP ANALYSIS: {os.path.basename(filepath)}")
    print(f"{'='*80}\n")
    
    print(f"Total Paragraphs: {len(doc.paragraphs)}")
    print(f"Total Sections: {len(doc.sections)}")
    print(f"Total Tables: {len(doc.tables)}")
    
    # Find all headings
    print("\n### ALL HEADINGS ###")
    for i, para in enumerate(doc.paragraphs):
        if para.style.name.startswith('Heading'):
            print(f"  [{i:3d}] {para.style.name:15s} | {para.text[:80]}")
    
    # Find TOC, LOF, LOT
    print("\n### SPECIAL SECTIONS (TOC/LOF/LOT) ###")
    for i, para in enumerate(doc.paragraphs):
        text_lower = para.text.lower()
        if 'table of contents' in text_lower or 'list of figures' in text_lower or 'list of tables' in text_lower:
            print(f"  [{i:3d}] {para.style.name:15s} | {para.text[:80]}")
    
    # Check for fields (TOC fields, page fields, etc.)
    print("\n### FIELDS DETECTED ###")
    field_count = 0
    for i, para in enumerate(doc.paragraphs[:100]):  # Check first 100 paragraphs
        for run in para.runs:
            if run._element.xpath('.//w:fldChar'):
                field_count += 1
                print(f"  Para {i}: Field detected in '{para.text[:50]}'")
                break
    if field_count == 0:
        print("  No fields detected in first 100 paragraphs")
    
    # Section breaks and page numbering
    print("\n### SECTION DETAILS ###")
    from docx.oxml.ns import qn
    for i, section in enumerate(doc.sections):
        print(f"\nSection {i}:")
        print(f"  Start type: {section.start_type}")
        
        # Page numbering
        sectPr = section._sectPr
        pgNumType = sectPr.find(qn('w:pgNumType'))
        if pgNumType is not None:
            fmt = pgNumType.get(qn('w:fmt'))
            start = pgNumType.get(qn('w:start'))
            print(f"  Page numbering: format={fmt}, start={start}")
        else:
            print(f"  Page numbering: Not set")
        
        # Headers/Footers
        if section.header:
            header_text = ' '.join([p.text for p in section.header.paragraphs])
            if header_text.strip():
                print(f"  Header: {header_text[:60]}")
        
        if section.footer:
            footer_text = ' '.join([p.text for p in section.footer.paragraphs])
            if footer_text.strip():
                print(f"  Footer: {footer_text[:60]}")
    
    # Check for chapter markers
    print("\n### CHAPTER MARKERS ###")
    for i, para in enumerate(doc.paragraphs):
        text_upper = para.text.upper()
        if 'CHAPTER' in text_upper and len(para.text) < 100:
            print(f"  [{i:3d}] {para.style.name:15s} | {para.text}")
    
    return doc

# Analyze both documents
base_path = r"C:\Users\Cephas\Documents\Programming\Student-Report-Formatter\examples"

print("\n" + "="*80)
print("CURRENT OUTPUT (result.docx)")
print("="*80)
result_doc = deep_analyze(os.path.join(base_path, "result.docx"))

print("\n\n" + "="*80)
print("EXPECTED OUTPUT (sample_edited.docx)")
print("="*80)
expected_doc = deep_analyze(os.path.join(base_path, "sample_edited.docx"))

print("\n\n" + "="*80)
print("ORIGINAL INPUT (sample.docx)")
print("="*80)
original_doc = deep_analyze(os.path.join(base_path, "sample.docx"))
