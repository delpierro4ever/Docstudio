"""
Deep dive: compare sample.docx and sample_edited.docx paragraph by paragraph
to understand what content was added/changed.
"""
from docx import Document
import os

EXAMPLES_DIR = os.path.join(os.path.dirname(__file__), "..", "examples")

def dump_all_paragraphs(filepath):
    doc = Document(filepath)
    print(f"\n{'='*80}")
    print(f"ALL PARAGRAPHS: {os.path.basename(filepath)} ({len(doc.paragraphs)} total)")
    print(f"{'='*80}")
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        style = p.style.name
        marker = ""
        if text:
            marker = f"[{style:20s}] {text[:150]}"
        else:
            marker = f"[{style:20s}] (empty)"
        print(f"  [{i:3d}] {marker}")

# Input file
dump_all_paragraphs(os.path.join(EXAMPLES_DIR, "sample.docx"))

# Expected output
dump_all_paragraphs(os.path.join(EXAMPLES_DIR, "sample_edited.docx"))
