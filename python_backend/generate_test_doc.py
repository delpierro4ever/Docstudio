from docx import Document

def create_test_doc(filename):
    doc = Document()
    
    # Cover content
    doc.add_paragraph("UNIVERSITY OF NOWHERE")
    doc.add_paragraph("THESIS TITLE")
    doc.add_page_break()
    
    # Title content
    doc.add_paragraph("THESIS TITLE AGAIN")
    doc.add_page_break()
    
    # Prelims
    doc.add_heading("DECLARATION", level=1)
    doc.add_paragraph("I declare things.")
    
    doc.add_heading("ABSTRACT", level=1)
    doc.add_paragraph("This is the abstract.")
    
    # Main Content
    doc.add_heading("CHAPTER 1", level=1)
    doc.add_heading("INTRODUCTION", level=1)
    doc.add_paragraph("This is the start of the thesis.")
    
    doc.save(filename)
    print(f"Created {filename}")

if __name__ == "__main__":
    create_test_doc("synthetic_test.docx")
