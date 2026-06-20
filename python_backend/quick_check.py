"""
Quick analysis of the latest output
"""
from docx import Document
from docx.oxml.ns import qn
import os

output_file = "downloaded_c9cce4f1-49c3-465a-9b03-2ec10629a564.docx"

if os.path.exists(output_file):
    doc = Document(output_file)
    
    print(f"Total Paragraphs: {len(doc.paragraphs)}")
    print(f"Total Sections: {len(doc.sections)}")
    
    print("\n### First 40 Paragraphs with Styles ###")
    for i, p in enumerate(doc.paragraphs[:40]):
        text = p.text[:80] if p.text else ""
        print(f"  [{i:3d}] {p.style.name:20s} | {text}")
    
    print("\n### Section Info ###")
    for i, section in enumerate(doc.sections):
        sectPr = section._sectPr
        pgNumType = sectPr.find(qn('w:pgNumType'))
        if pgNumType is not None:
            fmt = pgNumType.get(qn('w:fmt'))
            start = pgNumType.get(qn('w:start'))
            print(f"Section {i}: fmt={fmt}, start={start}")
        else:
            print(f"Section {i}: No page numbering set")
else:
    print(f"File not found: {output_file}")
