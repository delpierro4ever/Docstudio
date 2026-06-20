import docx
import os

def read_docx(file_path):
    # Resolve absolute path
    abs_path = os.path.abspath(file_path)
    if not os.path.exists(abs_path):
        print(f"File not found: {abs_path}")
        return

    try:
        doc = docx.Document(abs_path)
        print(f"Successfully opened: {os.path.basename(abs_path)}")
        print("-" * 30)
        
        # Print first 5 non-empty paragraphs
        count = 0
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                print(text)
                count += 1
            if count >= 5:
                break
                
        print("-" * 30)
        print(f"Total paragraphs: {len(doc.paragraphs)}")
        
    except Exception as e:
        print(f"Error reading file: {e}")

if __name__ == "__main__":
    # Pointing to the sample.docx in examples folder
    target_file = os.path.join(os.path.dirname(os.path.dirname(__file__)), "examples", "sample.docx")
    read_docx(target_file)
