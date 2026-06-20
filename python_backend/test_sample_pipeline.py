"""
Test the pipeline on examples/sample.docx and compare output to expected.
"""
import sys
import os
sys.path.insert(0, os.path.dirname(__file__))

from docx import Document
from app.extraction.doc_parser import extract_paragraph_stream
from app.intelligence.llm_client import analyze_structure_with_llm
from app.formatting.rebuilder import rebuild_document

def main():
    input_path = os.path.join(os.path.dirname(__file__), "..", "examples", "sample.docx")
    output_path = os.path.join(os.path.dirname(__file__), "..", "examples", "test_output.docx")
    
    print(f"Input: {input_path}")
    print(f"Output: {output_path}")
    
    # Load
    doc = Document(input_path)
    print(f"Loaded {len(doc.paragraphs)} paragraphs")
    
    # Extract
    stream = extract_paragraph_stream(doc)
    print(f"Extracted {len(stream)} paragraph metadata")
    
    # Analyze
    metadata = analyze_structure_with_llm(stream)
    print(f"Analysis: doc_type={metadata.doc_type}, main_start={metadata.anchors.main_content_start_idx}")
    
    # Rebuild
    new_doc = rebuild_document(doc, metadata)
    
    # Save
    new_doc.save(output_path)
    print(f"\nSaved to {output_path}")
    
    # Quick comparison
    print(f"\n{'='*60}")
    print("OUTPUT STRUCTURE:")
    print(f"{'='*60}")
    print(f"Paragraphs: {len(new_doc.paragraphs)}")
    print(f"Sections: {len(new_doc.sections)}")
    
    print(f"\nKey paragraphs:")
    for i, p in enumerate(new_doc.paragraphs):
        text = p.text.strip()
        if text:
            print(f"  [{i:3d}] [{p.style.name:20s}] {text[:100]}")
    
    # Compare with expected
    expected_path = os.path.join(os.path.dirname(__file__), "..", "examples", "sample_edited.docx")
    if os.path.exists(expected_path):
        expected = Document(expected_path)
        print(f"\n{'='*60}")
        print("COMPARISON WITH EXPECTED:")
        print(f"{'='*60}")
        print(f"Paragraphs: output={len(new_doc.paragraphs)} vs expected={len(expected.paragraphs)}")
        print(f"Sections: output={len(new_doc.sections)} vs expected={len(expected.sections)}")

if __name__ == "__main__":
    main()
