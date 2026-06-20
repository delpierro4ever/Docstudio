
from docx import Document
from docx.shared import Pt
import sys
import os

# Add app to path
sys.path.append(os.path.join(os.path.dirname(__file__), '..'))

from app.formatting.rebuilder import rebuild_document
from app.schemas.thesis_structure import ThesisStructure, AnchorMap, ChapterItem, CaptionRules

def verify():
    # 1. Create Input Doc
    doc = Document()
    doc.add_paragraph("Title Page Content")
    doc.add_paragraph("Abstract Content")
    p_chap = doc.add_paragraph("Chapter 1 Introduction")
    p_body = doc.add_paragraph("This is some body text.")
    
    # 2. Mock Metadata
    # Indices: 0=Title, 1=Abstract, 2=Chapter 1, 3=Body
    metadata = ThesisStructure(
        doc_type="report",
        confidence=1.0,
        anchors=AnchorMap(
            cover_start_idx=0,
            main_content_start_idx=2
        ),
        section_order=["cover", "main"],
        chapter_map=[
            ChapterItem(chapter_no=1, start_idx=2, end_idx=3, title="Introduction")
        ],
        caption_rules=CaptionRules()
    )
    
    # 3. Compile
    print("Running rebuild_document...")
    new_doc = rebuild_document(doc, metadata)
    
    # 4. Verify
    print("\nVerifying...")
    failures = []
    
    # Check Heading 1 Size (Chapter 1)
    # The rebuilder should have applied Heading 1 to index 2
    p_chap_new = new_doc.paragraphs[2 + 5] # +5 because rebuilder adds SectionBreak, TOC header, TOC field, Break, LOF Header... ??
    # Wait, rebuilder adds pages BEFORE Chapter 1.
    # Let's find "Chapter 1 Introduction" by text
    target_p = None
    for p in new_doc.paragraphs:
        if "Chapter 1 Introduction" in p.text:
            target_p = p
            break
            
    if not target_p:
        failures.append("Could not find 'Chapter 1 Introduction' in output")
    else:
        # Check Style
        if target_p.style.name != "Heading 1":
             failures.append(f"Chapter style is {target_p.style.name}, expected Heading 1")
        
        # Check Font Size
        # Style level or Run level?
        # Formatter sets STYLE font size.
        h1_style = new_doc.styles['Heading 1']
        size = h1_style.font.size
        print(f"Heading 1 Style Size: {size} (Expected {16*12700} or 16pt)")
        if size != Pt(16):
            failures.append(f"Heading 1 Size is {size}, expected {Pt(16)}")
            
    # Check Normal Spacing
    normal_style = new_doc.styles['Normal']
    spacing = normal_style.paragraph_format.line_spacing
    print(f"Normal Style Line Spacing: {spacing} (Expected None/Single)")
    if spacing is not None and spacing != 1.0: # None means default (usually single), 1.0 is single
         # If it's 1.5, that's a fail
         if spacing == 1.5:
             failures.append("Normal style still has 1.5 line spacing")

    if failures:
        print("\nFAILED:")
        for f in failures:
            print(f"- {f}")
        sys.exit(1)
    else:
        print("\nSUCCESS: All checks passed!")

if __name__ == "__main__":
    verify()
